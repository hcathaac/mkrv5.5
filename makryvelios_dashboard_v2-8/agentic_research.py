"""Agentic Research Mode backend for Makryvelios v5.8.1.

The agent has two layers:
1) deterministic/offline planning, evidence indexing, analysis and drafting;
2) optional LLM synthesis invoked explicitly by the UI.

No external model is required for numerical execution or package generation.
"""
from __future__ import annotations

import html
import io
import json
import math
import re
import zipfile
from collections import Counter
from difflib import SequenceMatcher
from itertools import combinations
from dataclasses import dataclass, field
from typing import Any, Sequence, Callable

import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from matplotlib import pyplot as plt

from analytics_core import (
    categorical_summary,
    cluster_table,
    correlation_matrix,
    descriptive_statistics,
    fit_detailed_model,
    group_tests,
    pca_table,
    quality_summary,
    to_excel_bytes,
)
from research_chair import docx_bytes


@dataclass
class AgenticPlan:
    goal: str
    steps: list[dict[str, str]]
    warnings: list[str]
    mappings: dict[str, Any]


@dataclass
class AgenticRun:
    plan: AgenticPlan
    tables: dict[str, pd.DataFrame] = field(default_factory=dict)
    narratives: dict[str, str] = field(default_factory=dict)
    figures: dict[str, bytes] = field(default_factory=dict)
    interactive_html: dict[str, bytes] = field(default_factory=dict)
    manifest: dict[str, Any] = field(default_factory=dict)


def extract_source_evidence(pdf_pages: pd.DataFrame) -> pd.DataFrame:
    """Create a page-level literature/source index without inventing bibliography metadata."""
    if pdf_pages is None or pdf_pages.empty:
        return pd.DataFrame(columns=["document", "page", "evidence_type", "doi", "url", "year_candidates", "snippet"])
    doi_re = re.compile(r"10\.\d{4,9}/[-._;()/:A-Z0-9]+", re.I)
    url_re = re.compile(r"https?://[^\s<>\]\)]+", re.I)
    year_re = re.compile(r"\b(?:19|20)\d{2}\b")
    rows = []
    for row in pdf_pages.itertuples(index=False):
        text = str(getattr(row, "text", "") or "")
        if not text.strip():
            rows.append({"document": row.document, "page": row.page, "evidence_type": "image-only/empty-text page", "doi": "", "url": "", "year_candidates": "", "snippet": ""})
            continue
        dois = list(dict.fromkeys(m.group(0).rstrip(".,;") for m in doi_re.finditer(text)))
        urls = list(dict.fromkeys(m.group(0).rstrip(".,;") for m in url_re.finditer(text)))
        years = list(dict.fromkeys(year_re.findall(text)))
        # Extractive synopsis: first substantive sentence plus one methods/results-like sentence if available.
        sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", re.sub(r"\s+", " ", text)) if len(s.strip()) >= 45]
        candidates = []
        if sentences:
            candidates.append(sentences[0])
            methodish = next((s for s in sentences if re.search(r"\b(method|model|regression|sample|result|finding|estimate|data|analysis|criterion|optim|simulation)\b", s, re.I)), None)
            if methodish and methodish not in candidates:
                candidates.append(methodish)
        snippet = " ".join(candidates)[:1200]
        rows.append({
            "document": row.document,
            "page": int(row.page),
            "evidence_type": "text-extracted page",
            "doi": "; ".join(dois[:5]),
            "url": "; ".join(urls[:5]),
            "year_candidates": "; ".join(years[:10]),
            "snippet": snippet,
        })
    return pd.DataFrame(rows)


def literature_key_terms(pdf_pages: pd.DataFrame, top_n: int = 30) -> pd.DataFrame:
    if pdf_pages is None or pdf_pages.empty:
        return pd.DataFrame(columns=["term", "count"])
    stop = {
        "that","this","with","from","were","have","has","had","which","their","there","these","those","into","between","using","used","than","also","such","more","most","through","within","where","when","what","would","could","should","study","paper","results","result","data","analysis","model","models","method","methods","table","figure","et","al","and","the","for","are","was","not","but","can","may","our","we","of","to","in","a","an","on","as","by","is","be","or","at","it"
    }
    tokens = []
    for text in pdf_pages.text.fillna("").astype(str):
        tokens.extend(re.findall(r"\b[A-Za-z][A-Za-z\-]{3,}\b", text.lower()))
    counts = Counter(t for t in tokens if t not in stop and not t.isdigit())
    return pd.DataFrame(counts.most_common(int(top_n)), columns=["term", "count"])


def generate_research_questions(df: pd.DataFrame, pdf_pages: pd.DataFrame | None = None, limit: int = 150) -> pd.DataFrame:
    """Generate a data-aware deterministic RQ bank from actual patterns, not raw cartesian templates."""
    limit = min(max(1, int(limit)), 200)
    numeric = list(df.select_dtypes(include=np.number).columns)
    categorical = [c for c in df.columns if c not in numeric and 2 <= df[c].nunique(dropna=True) <= 80]
    time_cols = [c for c in df.columns if re.search(r"year|date|time|έτος|χρον", str(c), re.I)]
    geo_cols = [c for c in df.columns if re.search(r"region|nuts|municip|geograph|περιφέρ|δήμ", str(c), re.I)]
    terms = literature_key_terms(pdf_pages, top_n=16).term.tolist() if pdf_pages is not None else []
    rows: list[dict[str, Any]] = []
    seen = set()

    def add(question: str, family: str, variables: Sequence[str], priority: int = 2, rationale: str = "", source_basis: str = ""):
        q = re.sub(r"\s+", " ", question).strip()
        key = _normalise_query(q)
        if q and key not in seen and len(rows) < limit:
            seen.add(key)
            rows.append({
                "rq_id": f"RQ{len(rows)+1:03d}",
                "research_question": q,
                "method_family": family,
                "variables": "; ".join(map(str, variables)),
                "priority": priority,
                "rationale": rationale,
                "source_basis": source_basis,
            })

    # Literature-only mode remains useful with no spreadsheet.
    if not numeric and terms:
        for term in terms:
            add(f"How is '{term}' defined and operationalised across the uploaded sources, and where do the definitions materially diverge?", "Literature synthesis", [], 1, "Frequent uploaded-PDF term", term)
            add(f"Which empirical designs and measurement strategies are used for '{term}', and which limitations recur across the uploaded sources?", "Methods review", [], 1, "Frequent uploaded-PDF term", term)
            add(f"What unresolved contradiction or research gap concerning '{term}' is directly visible in the uploaded evidence?", "Research gap", [], 2, "Frequent uploaded-PDF term", term)

    # Data-aware univariate priorities: missingness, variation and scale.
    profile = []
    for v in numeric[:120]:
        s = pd.to_numeric(df[v], errors="coerce")
        miss = float(s.isna().mean())
        std = float(s.std()) if s.notna().sum() > 1 else 0.0
        mean = float(s.mean()) if s.notna().any() else 0.0
        cv = abs(std / mean) if mean not in {0.0, -0.0} and np.isfinite(mean) else np.nan
        profile.append((miss, cv if np.isfinite(cv) else -1, std, v))
    for miss, cv, std, v in sorted(profile, reverse=True)[:25]:
        why = f"observed missingness={miss:.1%}; SD={std:.4g}" + (f"; |CV|={cv:.3f}" if cv >= 0 else "")
        add(f"Does the observed distribution of {v} contain enough dispersion and complete information for substantive modelling, and how sensitive are conclusions to its missingness and extremes?", "Descriptive / data quality", [v], 1, why, "Active dataset profile")

    # Rank actual pairwise patterns; do not enumerate arbitrary pairs first.
    ranked_pairs = []
    if len(numeric) >= 2:
        use = numeric[:60]
        try:
            pear = df[use].corr(method="pearson", numeric_only=True)
            spear = df[use].corr(method="spearman", numeric_only=True)
            for i, a in enumerate(use):
                for b in use[i + 1:]:
                    r = pear.loc[a, b]
                    rs = spear.loc[a, b]
                    if pd.notna(r):
                        stability = abs(float(r) - float(rs)) if pd.notna(rs) else np.nan
                        ranked_pairs.append((abs(float(r)), a, b, float(r), float(rs) if pd.notna(rs) else np.nan, stability))
        except Exception:
            ranked_pairs = []
    for abs_r, a, b, r, rs, stability in sorted(ranked_pairs, reverse=True)[:45]:
        stable_text = f"Pearson r={r:.3f}; Spearman ρ={rs:.3f}" if np.isfinite(rs) else f"Pearson r={r:.3f}"
        add(f"Why do {a} and {b} show one of the strongest observed relationships in the active data ({stable_text}), and does that relationship remain after relevant covariate adjustment?", "Association / adjusted econometrics", [a, b], 1, "Ranked from observed correlation structure", stable_text)
        if np.isfinite(stability) and stability > 0.15:
            add(f"The Pearson and rank-based association between {a} and {b} diverges materially. Is that discrepancy driven by outliers, non-linearity or subgroup structure?", "Robustness / non-linearity", [a, b], 1, "Pearson-Spearman divergence", stable_text)
        if len(rows) >= limit:
            break

    # Predictive/Bayesian questions use actual top partners rather than first columns.
    partners: dict[str, list[str]] = {v: [] for v in numeric}
    for _, a, b, *_ in sorted(ranked_pairs, reverse=True):
        if b not in partners[a] and len(partners[a]) < 6:
            partners[a].append(b)
        if a not in partners[b] and len(partners[b]) < 6:
            partners[b].append(a)
    for y in numeric[:30]:
        preds = partners.get(y, [])[:5]
        if preds:
            add(f"How well can {y} be predicted out of sample from its strongest observed covariates ({', '.join(preds)}), and which variables drive the predictions under SHAP/permutation explanation?", "Predictive + explainable ML", [y, *preds], 2, "Predictors chosen from observed association ranking", "Active dataset")
            add(f"What is the posterior uncertainty around the multivariable relationship between {y} and {', '.join(preds[:3])}, including posterior predictive fit rather than only point estimates?", "Bayesian modelling", [y, *preds[:3]], 2, "Variables chosen from observed association ranking", "Active dataset")

    # Group heterogeneity selects groups with actual usable sample sizes.
    for g in categorical[:15]:
        counts = df[g].value_counts(dropna=True)
        if len(counts) < 2 or counts.iloc[:2].min() < 5:
            continue
        for y in numeric[:15]:
            add(f"Does {y} differ substantively across {g} categories, which categories drive the heterogeneity, and is the effect robust to unequal variances and non-normality?", "Group comparison", [y, g], 2, f"{g} has {len(counts)} observed categories", "Active dataset")
            if len(rows) >= limit:
                break
        if len(rows) >= limit:
            break

    for t in time_cols[:5]:
        for y in numeric[:18]:
            add(f"What temporal pattern does {y} exhibit over {t}, and does the apparent trend survive alternative functional forms and structural-break checks?", "Longitudinal / time series", [y, t], 2, "Detected time-like field", str(t))
            if len(rows) >= limit:
                break
    for g in geo_cols[:5]:
        for y in numeric[:18]:
            add(f"Where are the spatial concentrations of high and low {y} across {g}, and do Moran/LISA diagnostics indicate clustering beyond a simple regional ranking?", "Spatial / GIS", [y, g], 2, "Detected geography field", str(g))
            if len(rows) >= limit:
                break

    # Causal design candidates are explicit design questions only.
    binary = [c for c in df.columns if df[c].nunique(dropna=True) == 2]
    for treatment in binary[:10]:
        for y in numeric[:12]:
            if treatment == y:
                continue
            add(f"If {treatment} can be defended as a treatment/exposure, what is its AIPW average treatment effect on {y}, is propensity-score overlap adequate, and which confounders are required for identification?", "Causal inference candidate", [treatment, y], 3, "Binary exposure candidate detected; causal interpretation requires assumptions", "Active dataset")
            if len(rows) >= limit:
                break

    # Multi-objective questions prioritise high-variance / likely score-budget variables.
    objective_vars = [v for _, _, _, v in sorted(profile, reverse=True) if re.search(r"score|benefit|budget|cost|fund|impact|criterion|absorp|score|δαπαν|προϋ", str(v), re.I)]
    objective_vars += [v for v in numeric if v not in objective_vars]
    for a, b in combinations(objective_vars[:12], 2):
        add(f"What is the Pareto frontier when optimising {a} and {b} simultaneously, which alternatives are dominated, and how stable is the frontier under parameter uncertainty?", "Pareto / robust optimisation", [a, b], 2, "Candidate objectives selected from semantic field names and observed variability", "Active dataset")
        if len(rows) >= limit:
            break

    # Literature-linked empirical questions reference the actual extracted themes.
    for term in terms[:16]:
        for y in numeric[:10]:
            add(f"Do the observed patterns in {y} support, qualify or contradict the uploaded literature's treatment of '{term}', and which page-level evidence should anchor that comparison?", "Literature-linked empirical synthesis", [y], 2, "Frequent uploaded-PDF term", term)
            if len(rows) >= limit:
                break

    # Fill remaining slots with robustness questions using ranked actual pairs.
    cycle = 0
    pair_source = sorted(ranked_pairs, reverse=True) or [(0.0, a, b, np.nan, np.nan, np.nan) for a, b in combinations(numeric[:20], 2)]
    while len(rows) < limit and pair_source:
        _, a, b, r, rs, _ = pair_source[cycle % len(pair_source)]
        add(f"How sensitive is the observed {a}–{b} relationship to missing-data handling, outlier treatment, subgroup composition, robust covariance and alternative model specifications?", "Robustness / sensitivity", [a, b], 3, "Pair selected from observed relationship ranking", f"r={r:.3f}" if np.isfinite(r) else "Active dataset")
        cycle += 1
        if cycle > limit * 5:
            break
    return pd.DataFrame(rows)

def build_agentic_plan(
    df: pd.DataFrame,
    *,
    goal: str,
    outcome: str | None,
    predictors: Sequence[str],
    group: str | None = None,
    time_column: str | None = None,
    region: str | None = None,
    pdf_pages: pd.DataFrame | None = None,
) -> AgenticPlan:
    predictors = [p for p in predictors if p in df.columns and p != outcome]
    steps = []
    if not df.empty:
        steps.append({"step": str(len(steps)+1), "action": "Audit active data", "engine": "deterministic", "output": "quality and missingness tables"})
    if pdf_pages is not None and not pdf_pages.empty:
        steps.append({"step": str(len(steps)+1), "action": "Index uploaded literature PDFs", "engine": "local PDF extraction", "output": "page-level evidence/source index"})
    steps.append({"step": str(len(steps)+1), "action": "Generate and rank research questions", "engine": "deterministic question generator", "output": "up to 150 RQs"})
    if not df.empty:
        steps.append({"step": str(len(steps)+1), "action": "Run descriptive statistics and correlation screen", "engine": "pandas/scipy", "output": "descriptive and association tables"})
    if outcome and predictors:
        steps.append({"step": "5", "action": f"Estimate primary OLS model for {outcome}", "engine": "statsmodels HC3", "output": "coefficients, diagnostics and fitted values"})
    if group and outcome:
        steps.append({"step": "6", "action": f"Test heterogeneity of {outcome} across {group}", "engine": "Welch/Mann-Whitney or ANOVA/Kruskal", "output": "tests + effect sizes"})
    if len(predictors) >= 2:
        steps.append({"step": "7", "action": "Run PCA and unsupervised segmentation", "engine": "scikit-learn", "output": "loadings, explained variance and cluster profiles"})
    if time_column:
        steps.append({"step": "8", "action": f"Preserve {time_column} as temporal routing information", "engine": "method router", "output": "time-aware recommendations"})
    if region:
        steps.append({"step": "9", "action": f"Preserve {region} for GIS/spatial routing", "engine": "method router", "output": "spatial-analysis recommendation"})
    steps.extend([
        {"step": str(len(steps)+1), "action": "Synthesize evidence-grounded discussion and conclusions", "engine": "offline deterministic narrative", "output": "draft discussion + conclusion"},
        {"step": str(len(steps)+2), "action": "Build submission package", "engine": "local export system", "output": "DOCX + XLSX + JSON + interactive HTML + publication graphics"},
    ])
    warnings = [
        "The agent never infers causality from an observational table without an explicit causal design.",
        "Extracted PDF text is research evidence, not automatically verified bibliographic metadata or quotation-ready text.",
        "The package is a near-submission draft and still requires human verification of claims, references, formatting and institutional requirements.",
        "External LLM use is optional; numerical analysis and the offline package remain available without an API key.",
    ]
    mappings = {"outcome": outcome, "predictors": list(predictors), "group": group, "time_column": time_column, "region": region, "pdf_documents": sorted(pdf_pages.document.unique().tolist()) if pdf_pages is not None and not pdf_pages.empty else []}
    return AgenticPlan(goal=goal.strip(), steps=steps, warnings=warnings, mappings=mappings)


def _top_correlations(corr: pd.DataFrame, pvals: pd.DataFrame, n: int = 30) -> pd.DataFrame:
    rows = []
    cols = list(corr.columns)
    for i, a in enumerate(cols):
        for b in cols[i+1:]:
            r = corr.loc[a, b]
            p = pvals.loc[a, b]
            if pd.notna(r):
                rows.append({"variable_1": a, "variable_2": b, "correlation": float(r), "abs_correlation": abs(float(r)), "p_value": float(p) if pd.notna(p) else np.nan})
    return pd.DataFrame(rows).sort_values("abs_correlation", ascending=False).head(n) if rows else pd.DataFrame(columns=["variable_1","variable_2","correlation","abs_correlation","p_value"])


def _correlation_figure_bytes(corr: pd.DataFrame) -> tuple[dict[str, bytes], dict[str, bytes]]:
    if corr.empty:
        return {}, {}
    fig, ax = plt.subplots(figsize=(10, 8), constrained_layout=True)
    image = ax.imshow(corr.to_numpy(float), vmin=-1, vmax=1, cmap="coolwarm", aspect="auto")
    ax.set_xticks(range(len(corr.columns))); ax.set_xticklabels(corr.columns, rotation=90, fontsize=7)
    ax.set_yticks(range(len(corr.index))); ax.set_yticklabels(corr.index, fontsize=7)
    ax.set_title("Correlation matrix", loc="left", fontsize=14, fontweight="bold")
    fig.colorbar(image, ax=ax, fraction=.035, pad=.02)
    static = {}
    for fmt in ["png", "svg", "pdf"]:
        buf = io.BytesIO()
        fig.savefig(buf, format=fmt, dpi=600 if fmt == "png" else None, bbox_inches="tight", facecolor="white")
        static[f"correlation_matrix.{fmt}"] = buf.getvalue()
    plt.close(fig)
    try:
        import plotly.express as px
        pfig = px.imshow(corr, zmin=-1, zmax=1, color_continuous_scale="RdBu_r", title="Interactive correlation matrix", aspect="auto")
        interactive = {"correlation_matrix.html": pfig.to_html(full_html=True, include_plotlyjs=True).encode("utf-8")}
    except Exception:
        interactive = {}
    return static, interactive


def _ols_figure_bytes(coef: pd.DataFrame) -> tuple[dict[str, bytes], dict[str, bytes]]:
    if coef is None or coef.empty or not {"term", "coefficient", "ci_95_low", "ci_95_high"}.issubset(coef.columns):
        return {}, {}
    d = coef[~coef.term.astype(str).str.lower().isin(["const", "intercept"])].copy().head(30)
    if d.empty:
        return {}, {}
    d = d.sort_values("coefficient")
    fig, ax = plt.subplots(figsize=(9, max(4, .36 * len(d) + 1.5)), constrained_layout=True)
    y = np.arange(len(d))
    ax.errorbar(d.coefficient, y, xerr=[d.coefficient - d.ci_95_low, d.ci_95_high - d.coefficient], fmt="o", capsize=3)
    ax.axvline(0, linewidth=1, linestyle="--")
    ax.set_yticks(y); ax.set_yticklabels(d.term, fontsize=8)
    ax.set_xlabel("Coefficient (95% CI)"); ax.set_title("Primary OLS coefficient estimates", loc="left", fontsize=14, fontweight="bold")
    ax.grid(axis="x", alpha=.2)
    static = {}
    for fmt in ["png", "svg", "pdf"]:
        buf = io.BytesIO(); fig.savefig(buf, format=fmt, dpi=600 if fmt == "png" else None, bbox_inches="tight", facecolor="white"); static[f"ols_coefficients.{fmt}"] = buf.getvalue()
    plt.close(fig)
    try:
        import plotly.express as px
        pfig = px.scatter(d, x="coefficient", y="term", error_x=d["ci_95_high"]-d["coefficient"], error_x_minus=d["coefficient"]-d["ci_95_low"], title="Primary OLS coefficient estimates")
        pfig.add_vline(x=0, line_dash="dash")
        interactive = {"ols_coefficients.html": pfig.to_html(full_html=True, include_plotlyjs=True).encode("utf-8")}
    except Exception:
        interactive = {}
    return static, interactive


def _normalise_query(text: str) -> str:
    return re.sub(r"[^a-z0-9α-ωάέήίόύώϊϋΐΰ_\- ]+", " ", str(text).casefold()).strip()


def _run_evidence_chunks(run: AgenticRun, max_rows_per_table: int = 60) -> list[dict[str, str]]:
    """Build searchable evidence chunks from the actual computed run."""
    chunks: list[dict[str, str]] = []
    for name, table in run.tables.items():
        if not isinstance(table, pd.DataFrame) or table.empty:
            continue
        view = table.head(max_rows_per_table)
        for idx, row in view.iterrows():
            bits = []
            for col, value in row.items():
                if pd.isna(value):
                    continue
                bits.append(f"{col}={value}")
            chunks.append({"source": name, "locator": str(idx), "text": f"{name}: " + "; ".join(bits)})
    for key, value in run.narratives.items():
        if value:
            chunks.append({"source": f"Narrative:{key}", "locator": "", "text": str(value)})
    return chunks


def semantic_retrieve_run(run: AgenticRun, question: str, top_k: int = 8) -> list[dict[str, str]]:
    chunks = _run_evidence_chunks(run)
    if not chunks or not str(question).strip():
        return []
    corpus = [c["text"] for c in chunks]
    try:
        vec = TfidfVectorizer(ngram_range=(1, 2), min_df=1, strip_accents="unicode")
        mat = vec.fit_transform(corpus + [str(question)])
        scores = cosine_similarity(mat[-1], mat[:-1]).ravel()
        order = np.argsort(scores)[::-1]
        out = []
        for i in order:
            if scores[i] <= 0 and out:
                break
            item = dict(chunks[int(i)])
            item["similarity"] = f"{float(scores[i]):.4f}"
            out.append(item)
            if len(out) >= int(top_k):
                break
        return out
    except Exception:
        q = _normalise_query(question)
        terms = set(q.split())
        scored = []
        for c in chunks:
            txt = _normalise_query(c["text"])
            overlap = sum(1 for t in terms if t and t in txt)
            if overlap:
                scored.append((overlap, c))
        return [dict(c, similarity=str(score)) for score, c in sorted(scored, key=lambda x: x[0], reverse=True)[:top_k]]


def _query_intent(question: str) -> str:
    q = _normalise_query(question)
    examples = {
        "weakest": [
            "weakest finding least convincing result what cannot conclude uncertain insignificant not significant caveat",
            "ποιο ειναι το πιο αδυναμο ευρημα τι δεν μπορω να συμπερανω αβεβαιο μη σημαντικο",
        ],
        "strongest": [
            "strongest finding most important key result strongest evidence main finding",
            "ισχυροτερο ευρημα σημαντικοτερο αποτελεσμα",
        ],
        "causality": [
            "causal cause effect can i say causes treatment effect causal inference",
            "αιτιωδης προκαλει επιδραση αιτιοτητα",
        ],
        "model": [
            "regression ols coefficient model fit predictor significance diagnostics",
            "παλινδρομηση μοντελο συντελεστης",
        ],
        "literature": [
            "literature pdf source citation evidence paper bibliography what do sources say",
            "βιβλιογραφια πηγες αρθρα pdf",
        ],
        "research_questions": [
            "research questions hypotheses rq questions generate",
            "ερευνητικα ερωτηματα υποθεσεις",
        ],
        "next": [
            "what next further analysis additional test robustness next step",
            "τι αλλο επομενο περαιτερω αναλυση",
        ],
        "conclusion": [
            "conclusion conclude overall meaning what can safely conclude",
            "συμπερασμα τι μπορω να πω",
        ],
        "compare": [
            "compare difference versus which is better stronger weaker between",
            "συγκρινε διαφορα μεταξυ",
        ],
        "explain": [
            "why explain interpret what does this mean meaning",
            "γιατι εξηγησε τι σημαινει ερμηνεια",
        ],
    }
    texts, labels = [], []
    for label, vals in examples.items():
        for value in vals:
            texts.append(value)
            labels.append(label)
    try:
        vec = TfidfVectorizer(ngram_range=(1, 2), strip_accents="unicode")
        mat = vec.fit_transform(texts + [q])
        scores = cosine_similarity(mat[-1], mat[:-1]).ravel()
        if float(scores.max()) >= 0.08:
            return labels[int(scores.argmax())]
    except Exception:
        pass
    return "open"


def _format_p(value: Any) -> str:
    try:
        return f"{float(value):.4g}"
    except Exception:
        return str(value)


def _specific_weakest_finding(run: AgenticRun) -> str:
    coef = run.tables.get("OLS coefficients", pd.DataFrame())
    if isinstance(coef, pd.DataFrame) and not coef.empty and "p_value" in coef:
        d = coef.copy()
        if "term" in d:
            d = d[~d.term.astype(str).str.lower().isin(["const", "intercept"])]
        d["p_num"] = pd.to_numeric(d["p_value"], errors="coerce")
        d = d.dropna(subset=["p_num"]).sort_values("p_num", ascending=False)
        if not d.empty:
            r = d.iloc[0]
            ci = ""
            crosses = None
            if {"ci_95_low", "ci_95_high"}.issubset(d.columns):
                lo, hi = float(r.ci_95_low), float(r.ci_95_high)
                ci = f", 95% CI [{lo:.4g}, {hi:.4g}]"
                crosses = lo <= 0 <= hi
            verdict = "The interval crosses zero, so the adjusted association is not clearly distinguished from the null." if crosses else "It is the least statistically supported coefficient among the configured predictors, even though its interval may not cross zero."
            return f"The weakest substantive finding in the configured primary OLS model is **{r.term}**: β={float(r.coefficient):.4g}{ci}, p={_format_p(r.p_num)}. {verdict}\n\n**What you cannot safely conclude:** do not present this predictor as a robust independent relationship if the interval includes zero or diagnostics/specification do not support it. The OLS model is also not, by itself, evidence that the predictor causes the outcome."
    gt = run.tables.get("Group tests", pd.DataFrame())
    if isinstance(gt, pd.DataFrame) and not gt.empty:
        pcol = next((c for c in gt.columns if str(c).lower() in {"p_value", "p", "pvalue"}), None)
        if pcol:
            d = gt.copy()
            if "test" in d:
                d = d[~d["test"].astype(str).str.contains("levene|normality|shapiro", case=False, na=False)]
            d["p_num"] = pd.to_numeric(d[pcol], errors="coerce")
            d = d.dropna(subset=["p_num"]).sort_values("p_num", ascending=False)
            if not d.empty:
                r = d.iloc[0]
                label = "; ".join(f"{c}={r[c]}" for c in d.columns[:5] if c != "p_num")
                return f"The weakest substantive group-comparison result is {label}, p={_format_p(r.p_num)}. It should not be described as a reliable group difference unless the effect size and uncertainty support that interpretation."
    top = run.tables.get("Top correlations", pd.DataFrame())
    if isinstance(top, pd.DataFrame) and not top.empty:
        d = top.copy()
        d["p_num"] = pd.to_numeric(d.get("p_value"), errors="coerce")
        d["abs_r"] = pd.to_numeric(d.get("abs_correlation"), errors="coerce")
        d = d.sort_values(["p_num", "abs_r"], ascending=[False, True])
        r = d.iloc[0]
        return f"Within the retained correlation screen, the weakest pair is **{r.variable_1} ↔ {r.variable_2}**: r={float(r.correlation):.3f}, p={_format_p(r.p_value)}. This does not support a causal claim and may not support even a stable association if the p-value is weak."
    return "This run does not contain enough ranked inferential output to identify a weakest empirical finding. Map an outcome/predictors or run the relevant tests first."


def _specific_named_model_term(run: AgenticRun, question: str) -> str | None:
    coef = run.tables.get("OLS coefficients", pd.DataFrame())
    if not isinstance(coef, pd.DataFrame) or coef.empty or "term" not in coef:
        return None
    q = _normalise_query(question)
    matches = []
    for _, row in coef.iterrows():
        term = str(row.get("term", ""))
        if term.lower() in {"const", "intercept"}:
            continue
        nt = _normalise_query(term)
        if nt and (nt in q or SequenceMatcher(None, nt, q).ratio() > 0.72):
            matches.append(row)
    if not matches:
        return None
    r = matches[0]
    ci = ""
    if "ci_95_low" in coef and "ci_95_high" in coef and pd.notna(r.get("ci_95_low")) and pd.notna(r.get("ci_95_high")):
        ci = f", 95% CI [{float(r['ci_95_low']):.4g}, {float(r['ci_95_high']):.4g}]"
    return f"For **{r['term']}** in the configured primary OLS model: β={float(r['coefficient']):.4g}{ci}, p={_format_p(r.get('p_value'))}. This is the adjusted association conditional on the other mapped predictors; it is not automatically causal."

def _specific_strongest_finding(run: AgenticRun) -> str:
    coef = run.tables.get("OLS coefficients", pd.DataFrame())
    if isinstance(coef, pd.DataFrame) and not coef.empty and "p_value" in coef:
        d = coef.copy()
        if "term" in d:
            d = d[~d.term.astype(str).str.lower().isin(["const", "intercept"])]
        d["p_num"] = pd.to_numeric(d["p_value"], errors="coerce")
        d = d.dropna(subset=["p_num"]).sort_values("p_num")
        if not d.empty:
            r = d.iloc[0]
            ci = ""
            if {"ci_95_low", "ci_95_high"}.issubset(d.columns):
                ci = f", 95% CI [{float(r.ci_95_low):.4g}, {float(r.ci_95_high):.4g}]"
            return f"The strongest result in the configured primary OLS model is **{r.term}**: β={float(r.coefficient):.4g}{ci}, p={_format_p(r.p_num)}. This is the strongest conditional association in that model, not automatically a causal effect."
    top = run.tables.get("Top correlations", pd.DataFrame())
    if isinstance(top, pd.DataFrame) and not top.empty:
        r = top.iloc[0]
        return f"The strongest screened pairwise association is **{r.variable_1} ↔ {r.variable_2}**: r={float(r.correlation):.3f}, p={_format_p(r.p_value)}. It is a bivariate association screen, not a causal estimate."
    return run.narratives.get("discussion", "No ranked empirical finding is available.")


def _literature_answer(run: AgenticRun, question: str) -> str:
    ev = run.tables.get("Literature source evidence", pd.DataFrame())
    if not isinstance(ev, pd.DataFrame) or ev.empty:
        return "No literature PDFs were indexed in this run."
    corpus = []
    for _, row in ev.head(2000).iterrows():
        corpus.append(f"{row.get('document', '')} page {row.get('page', '')}: {row.get('snippet', '')} {row.get('doi', '')}")
    try:
        vec = TfidfVectorizer(ngram_range=(1, 2), strip_accents="unicode", stop_words="english")
        mat = vec.fit_transform(corpus + [question])
        scores = cosine_similarity(mat[-1], mat[:-1]).ravel()
        order = np.argsort(scores)[::-1][:5]
    except Exception:
        order = np.arange(min(5, len(ev)))
    bullets = []
    for idx in order:
        row = ev.iloc[int(idx)]
        snippet = str(row.get("snippet", "")).strip()[:420]
        if snippet:
            doi = f" DOI {row.get('doi')}" if str(row.get("doi", "")).strip() else ""
            bullets.append(f"- **{row.get('document')}**, p. {row.get('page')}{doi}: {snippet}")
    return "The most relevant passages in the uploaded PDF evidence are:\n" + "\n".join(bullets) + "\n\nThese are extractive source notes; verify quotations and final bibliographic metadata before submission."


def offline_agent_reply(run: AgenticRun, question: str, history: Sequence[dict[str, str]] | None = None) -> str:
    """Evidence-grounded deterministic conversation with semantic routing."""
    q = str(question).strip()
    if not q:
        return "Ask a specific question about this run, a variable, a model, a PDF source, a finding, a limitation or the next analysis."
    named_term = _specific_named_model_term(run, q)
    if named_term is not None and not any(token in _normalise_query(q) for token in ["strongest", "weakest", "ισχυροτερο", "αδυναμο"]):
        return named_term
    intent = _query_intent(q)
    if intent == "weakest":
        return _specific_weakest_finding(run)
    if intent == "strongest":
        return _specific_strongest_finding(run)
    if intent == "literature":
        return _literature_answer(run, q)
    if intent == "research_questions":
        rqs = run.tables.get("Research questions", pd.DataFrame())
        if rqs.empty:
            return "No research-question bank was generated."
        try:
            vec = TfidfVectorizer(ngram_range=(1, 2), strip_accents="unicode")
            texts = rqs.research_question.astype(str).tolist()
            mat = vec.fit_transform(texts + [q])
            scores = cosine_similarity(mat[-1], mat[:-1]).ravel()
            top = rqs.iloc[np.argsort(scores)[::-1][:10]]
        except Exception:
            top = rqs.head(10)
        return "Most relevant generated research questions:\n" + "\n".join(f"{r.rq_id}. {r.research_question} [{r.method_family}]" for r in top.itertuples(index=False))
    if intent == "model":
        coef = run.tables.get("OLS coefficients", pd.DataFrame())
        fit = run.tables.get("OLS fit", pd.DataFrame())
        diag = run.tables.get("OLS diagnostics", pd.DataFrame())
        if coef.empty:
            return "No primary OLS model was executed. Map a primary outcome and predictors, rebuild the plan and approve the run."
        nonconst = coef[~coef.term.astype(str).str.lower().isin(["const", "intercept"])].copy()
        nonconst["p_num"] = pd.to_numeric(nonconst.p_value, errors="coerce")
        nonconst = nonconst.sort_values("p_num").head(8)
        terms = "; ".join(f"{r.term}: β={float(r.coefficient):.4g}, p={_format_p(r.p_value)}" for r in nonconst.itertuples(index=False))
        ftxt = fit.head(1).to_dict("records")[0] if not fit.empty else {}
        dtxt = diag.head(8).to_dict("records") if not diag.empty else []
        return f"Primary model fit: {ftxt}.\n\nMost supported non-intercept terms: {terms}.\n\nDiagnostics available: {dtxt}."
    if intent == "causality":
        causal = [(n, t) for n, t in run.tables.items() if isinstance(t, pd.DataFrame) and not t.empty and re.search(r"causal|aipw|treatment|ate", str(n), re.I)]
        if not causal:
            return "This Agentic run contains association/prediction outputs but no explicit causal-effect table. Therefore it is **not safe to say that one variable causes another** from this run. Use the Causal Inference Lab with a defensible treatment, outcome, covariate set, overlap check and identification assumptions."
        name, table = causal[0]
        return f"A causal-design output is present in **{name}**. Its first rows are:\n{table.head(8).to_string(index=False)}\n\nThe causal interpretation is still conditional on the stated identification assumptions and overlap/diagnostic checks."
    if intent == "conclusion":
        return run.narratives.get("conclusion", "No conclusion draft is available.") + "\n\n" + _specific_strongest_finding(run)
    if intent == "next":
        return "Recommended next analyses are conditional on the present evidence:\n- inspect the weakest/unstable terms and diagnostics before adding complexity;\n- use causal AIPW only for a defensible treatment/outcome design;\n- use Bayesian modelling when posterior uncertainty is substantively useful;\n- use SHAP for predictive explanation, not causal attribution;\n- use Pareto/robust optimisation for competing portfolio objectives;\n- re-run this Agentic workflow after any data or specification change so the manuscript and manifest remain aligned."
    retrieved = semantic_retrieve_run(run, q, top_k=6)
    if retrieved:
        evidence = "\n".join(f"- [{r['source']}] {r['text'][:650]}" for r in retrieved)
        return f"I matched your question to the following computed/source evidence from this run:\n{evidence}\n\nIf you want freer synthesis across these rows, use Local AI or External AI in this same Agentic tab; the numerical evidence above remains the source of truth."
    return "I could not find run evidence that directly matches that question. Ask about a named variable/model/source, or run the corresponding analysis first."


def agent_context_text(run: AgenticRun, question: str = "", max_chars: int = 50000) -> str:
    retrieved = semantic_retrieve_run(run, question or run.plan.goal, top_k=18)
    chunks = [f"RESEARCH GOAL: {run.plan.goal}"]
    if retrieved:
        chunks.append("MOST RELEVANT COMPUTED EVIDENCE:\n" + "\n".join(f"[{r['source']}] {r['text']}" for r in retrieved))
    for name in ["Quality audit", "Top correlations", "OLS coefficients", "OLS fit", "OLS diagnostics", "Group tests", "Literature key terms"]:
        table = run.tables.get(name)
        if isinstance(table, pd.DataFrame) and not table.empty:
            chunks.append(f"[{name}]\n{table.head(25).to_csv(index=False)}")
    ev = run.tables.get("Literature source evidence")
    if isinstance(ev, pd.DataFrame) and not ev.empty:
        chunks.append("[LITERATURE EVIDENCE]\n" + ev.head(40).to_csv(index=False))
    chunks.append("[OFFLINE DISCUSSION]\n" + run.narratives.get("discussion", ""))
    return "\n\n".join(chunks)[:int(max_chars)]


def ollama_text_reply(prompt: str, model: str, endpoint: str = "http://127.0.0.1:11434", timeout: int = 180, temperature: float = 0.12) -> str:
    import requests
    response = requests.post(
        f"{endpoint.rstrip('/')}/api/generate",
        json={"model": model, "prompt": str(prompt), "stream": False, "options": {"temperature": float(temperature)}},
        timeout=timeout,
    )
    response.raise_for_status()
    text = str(response.json().get("response", "")).strip()
    if not text:
        raise RuntimeError("Local Ollama returned no response text.")
    return text


def ollama_agent_reply(run: AgenticRun, question: str, model: str, endpoint: str = "http://127.0.0.1:11434", history: Sequence[dict[str, str]] | None = None, timeout: int = 180) -> str:
    import requests
    context = agent_context_text(run, question, max_chars=55000)
    hist = "\n".join(f"{m.get('role', 'user').upper()}: {m.get('content', '')}" for m in (history or [])[-8:])
    prompt = f"""You are the local Makryvelios research agent. Answer the user's exact question, not a generic safety template. Use only the supplied computed evidence and uploaded-PDF evidence. Quote numerical values exactly. If the evidence does not answer the question, say what is missing. Distinguish descriptive association, adjusted association, prediction, optimisation and causal estimates. For literature evidence, cite document name and page when available. Answer in the same language as the user unless explicitly asked otherwise.

RECENT CONVERSATION
{hist}

RUN EVIDENCE
{context}

USER QUESTION
{question}
"""
    return ollama_text_reply(prompt, model, endpoint=endpoint, timeout=timeout, temperature=0.12)


SYNTHESIS_RESPONSE_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "abstract": {"type": "string"},
        "results": {"type": "string"},
        "discussion": {"type": "string"},
        "conclusion": {"type": "string"},
        "limitations": {"type": "string"},
        "key_findings": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "finding": {"type": "string"},
                    "evidence": {"type": "string"},
                    "strength": {"type": "string"},
                    "safe_inference": {"type": "string"},
                },
                "required": ["finding", "evidence", "strength", "safe_inference"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["abstract", "results", "discussion", "conclusion", "limitations", "key_findings"],
    "additionalProperties": False,
}


def _sectioned_synthesis_from_prose(raw: str) -> dict[str, Any]:
    """Recover a structured synthesis from sectioned prose from the *same* model response.

    This is not a provider/model fallback. It simply accepts Markdown/plain-text
    headings when a model ignores the requested JSON envelope.
    """
    text = str(raw or "").strip()
    if not text:
        return {}
    aliases = {
        "abstract": "abstract",
        "executive abstract": "abstract",
        "results": "results",
        "results synthesis": "results",
        "findings": "results",
        "discussion": "discussion",
        "conclusion": "conclusion",
        "conclusions": "conclusion",
        "limitations": "limitations",
        "limitations and caveats": "limitations",
        "key findings": "key_findings",
        "key findings and safe inference": "key_findings",
    }
    heading_re = re.compile(
        r"(?im)^\s*(?:#{1,6}\s*|\*\*|__)?"
        r"(abstract|executive abstract|results synthesis|results|findings|discussion|conclusions?|limitations(?: and caveats)?|key findings(?: and safe inference)?)"
        r"(?:\*\*|__)?\s*[:\-–—]?\s*$"
    )
    matches = list(heading_re.finditer(text))
    if not matches:
        # Accept compact labelled sections such as 'Abstract: ... Results: ...'.
        inline_re = re.compile(
            r"(?is)(?:^|\n)\s*(abstract|results|discussion|conclusion|limitations|key findings)\s*:\s*"
        )
        matches = list(inline_re.finditer(text))
        heading_re = inline_re
    if not matches:
        return {}
    sections: dict[str, str] = {}
    for i, m in enumerate(matches):
        label = re.sub(r"\s+", " ", m.group(1).strip().lower())
        key = aliases.get(label, label)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        if body:
            sections[key] = body
    if "key_findings" in sections:
        findings = []
        for line in sections["key_findings"].splitlines():
            clean = re.sub(r"^\s*(?:[-*•]|\d+[.)])\s*", "", line).strip()
            if not clean:
                continue
            findings.append({
                "finding": clean,
                "evidence": "See the sectioned model response and the computed evidence supplied to the synthesis pass.",
                "strength": "model-synthesised",
                "safe_inference": "Interpret only at the level supported by the deterministic run.",
            })
        sections["key_findings"] = findings
    else:
        sections["key_findings"] = []
    return sections


def _parse_synthesis_json(text: str) -> dict[str, Any]:
    raw = str(text or "").strip()
    if not raw:
        raise RuntimeError("The selected AI returned an empty synthesis response.")
    cleaned = re.sub(r"```(?:json|javascript|js)?", "", raw, flags=re.I).replace("```", "").strip()

    # 1) Strict whole-response JSON.
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, dict):
            return obj
    except Exception:
        pass

    # 2) Embedded JSON object in explanatory prose.
    decoder = json.JSONDecoder()
    for match in re.finditer(r"\{", cleaned):
        try:
            obj, _ = decoder.raw_decode(cleaned[match.start():])
            if isinstance(obj, dict):
                return obj
        except Exception:
            continue

    # 3) Sectioned prose from the same model response. This preserves useful
    #    synthesis while still requiring all manuscript sections before the run
    #    is marked as AI-refined.
    sectioned = _sectioned_synthesis_from_prose(cleaned)
    required = ["abstract", "results", "discussion", "conclusion", "limitations"]
    if sectioned and all(str(sectioned.get(k, "")).strip() for k in required):
        return sectioned

    raise RuntimeError(
        "The selected AI response could not be structured into Abstract, Results, Discussion, Conclusion and Limitations. "
        "No provider/model fallback was used; retry the same model or choose another one manually."
    )


def refine_run_with_ai(
    run: AgenticRun,
    reply_fn: Callable[[str], str],
    *,
    provider_label: str = "AI",
    context_profile: str = "Compact",
) -> AgenticRun:
    """Rewrite the draft sections from the *actual computed run* and PDF evidence.

    The deterministic tables remain the source of truth. The function never changes
    model outputs; it only adds an evidence-grounded synthesis layer used by the
    paper/report exports.
    """
    context = agent_context_text(
        run,
        "strongest and weakest findings exact model results literature evidence discussion conclusion limitations",
        max_chars=_context_profile_limits(context_profile)["synthesis_chars"],
    )
    prompt = f"""You are refining a research draft from a completed deterministic analytical run.

RESEARCH GOAL
{run.plan.goal}

COMPUTED RESULTS AND UPLOADED-PDF EVIDENCE
{context}

TASK
Create a specific, evidence-grounded synthesis. Do NOT write generic statements such as 'the strongest defensible conclusions are those supported by the tables'. Instead name the actual variables, coefficients, correlations, confidence intervals, p-values, diagnostics, groups, sources and pages that support each statement. Do not invent bibliographic facts or causal claims. If a requested section is not supported, state the exact missing model/design rather than filling it with boilerplate.

Return ONLY one JSON object with these keys:
- abstract: 180-300 words
- results: detailed results narrative grounded in exact computed values
- discussion: interpret the actual findings and relate them to retrieved PDF evidence by document/page where available
- conclusion: concise substantive conclusion that directly answers the research goal
- limitations: limitations specific to THIS run, variables, models and evidence
- key_findings: array of 5-12 objects, each with finding, evidence, strength, and safe_inference

Keep association, prediction, optimisation and causal inference explicitly distinct. Preserve every supplied number exactly.
"""
    parsed = _parse_synthesis_json(reply_fn(prompt))
    required = ["abstract", "results", "discussion", "conclusion", "limitations"]
    missing = [k for k in required if not str(parsed.get(k, "")).strip()]
    if missing:
        raise RuntimeError("AI synthesis omitted required section(s): " + ", ".join(missing))

    # Preserve deterministic drafts for audit/reproducibility.
    for key in ["discussion", "conclusion", "limitations"]:
        if key in run.narratives and f"offline_{key}" not in run.narratives:
            run.narratives[f"offline_{key}"] = run.narratives[key]
    run.narratives["abstract"] = str(parsed["abstract"]).strip()
    run.narratives["results"] = str(parsed["results"]).strip()
    run.narratives["discussion"] = str(parsed["discussion"]).strip()
    run.narratives["conclusion"] = str(parsed["conclusion"]).strip()
    run.narratives["limitations"] = str(parsed["limitations"]).strip()

    findings = parsed.get("key_findings", [])
    if isinstance(findings, list):
        clean = []
        for item in findings:
            if isinstance(item, dict) and item.get("finding"):
                clean.append({
                    "finding": str(item.get("finding", "")),
                    "evidence": str(item.get("evidence", "")),
                    "strength": str(item.get("strength", "")),
                    "safe_inference": str(item.get("safe_inference", "")),
                })
        if clean:
            run.tables["AI synthesis key findings"] = pd.DataFrame(clean)

    run.manifest["ai_synthesis"] = {
        "enabled": True,
        "provider": provider_label,
        "scope": "narrative synthesis only; deterministic numerical tables unchanged",
    }
    return run


def _parse_rq_json(text: str) -> list[dict[str, Any]]:
    """Parse AI-generated research questions defensively.

    Accepts strict JSON, fenced JSON, objects wrapping arrays, JSON embedded in
    surrounding prose, JSONL, and numbered/bulleted plain-text questions.  This
    function deliberately degrades gracefully because hosted/local models can
    ignore formatting instructions even when the semantic content is useful.
    """
    raw = str(text or "").strip()
    if not raw:
        return []
    raw = re.sub(r"```(?:json|javascript|js)?", "", raw, flags=re.I).replace("```", "").strip()

    candidates: list[Any] = []

    def unpack(obj: Any) -> None:
        if isinstance(obj, list):
            candidates.extend(obj)
        elif isinstance(obj, dict):
            for key in ("questions", "research_questions", "items", "data", "results"):
                value = obj.get(key)
                if isinstance(value, list):
                    candidates.extend(value)
                    return
            # A single question object is also valid.
            if obj.get("research_question") or obj.get("question"):
                candidates.append(obj)

    # 1) Whole-response JSON.
    try:
        unpack(json.loads(raw))
    except Exception:
        pass

    # 2) Locate an embedded JSON array/object, tolerating explanatory text.
    if not candidates:
        decoder = json.JSONDecoder()
        for match in re.finditer(r"[\[{]", raw):
            try:
                obj, _ = decoder.raw_decode(raw[match.start():])
                unpack(obj)
                if candidates:
                    break
            except Exception:
                continue

    # 3) JSONL / one-object-per-line.
    if not candidates:
        for line in raw.splitlines():
            line = line.strip().rstrip(",")
            if not line:
                continue
            try:
                unpack(json.loads(line))
            except Exception:
                continue

    # 4) Plain numbered/bulleted question list fallback.  Preserve useful model
    #    output instead of failing the entire batch merely because JSON was not
    #    respected.
    if not candidates:
        for line in raw.splitlines():
            line = re.sub(r"^\s*(?:[-*•]|\d{1,3}[.)]|RQ\s*\d{1,3}[:.)-]?)\s*", "", line, flags=re.I).strip()
            if not line:
                continue
            # Extract a question-looking segment from prose/list lines.
            if "?" in line:
                q = line[: line.find("?") + 1].strip(' "\'')
                if len(q) >= 18:
                    candidates.append({"research_question": q})

    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in candidates:
        if isinstance(item, str):
            item = {"research_question": item}
        if not isinstance(item, dict):
            continue
        q = str(item.get("research_question") or item.get("question") or item.get("rq") or "").strip()
        q = re.sub(r"\s+", " ", q).strip()
        if not q:
            continue
        key = _normalise_query(q)
        if not key or key in seen:
            continue
        seen.add(key)
        priority_raw = item.get("priority", 2)
        try:
            priority = int(float(priority_raw))
        except Exception:
            priority = 2
        priority = min(3, max(1, priority))
        variables = item.get("variables", "")
        if isinstance(variables, (list, tuple)):
            variables = "; ".join(map(str, variables))
        out.append({
            "research_question": q,
            "method_family": str(item.get("method_family") or item.get("method") or item.get("analysis") or "AI-grounded"),
            "variables": str(variables or ""),
            "priority": priority,
            "rationale": str(item.get("rationale") or item.get("why") or item.get("reason") or ""),
            "source_basis": str(item.get("source_basis") or item.get("evidence") or item.get("source") or ""),
        })
    return out


RQ_RESPONSE_SCHEMA: dict[str, Any] = {
    "type": "array",
    "items": {
        "type": "object",
        "properties": {
            "research_question": {"type": "string"},
            "method_family": {"type": "string"},
            "variables": {"type": "string"},
            "priority": {"type": "integer"},
            "rationale": {"type": "string"},
            "source_basis": {"type": "string"},
        },
        "required": ["research_question", "method_family", "variables", "priority", "rationale", "source_basis"],
        "additionalProperties": False,
    },
}

def _context_profile_limits(profile: str) -> dict[str, int]:
    """Transparent context limits for hosted/local AI calls.

    These limits only change how much evidence is *sent* to the selected model;
    they never change the deterministic analyses or silently switch provider/model.
    """
    key = str(profile or "Compact").strip().lower()
    if key.startswith("extended"):
        return {"schema": 60, "corr": 30, "literature": 18, "snippet_chars": 800, "existing": 35, "goal_chars": 2200, "synthesis_chars": 42000}
    if key.startswith("standard"):
        return {"schema": 32, "corr": 18, "literature": 10, "snippet_chars": 600, "existing": 22, "goal_chars": 1600, "synthesis_chars": 24000}
    return {"schema": 18, "corr": 10, "literature": 5, "snippet_chars": 420, "existing": 12, "goal_chars": 1200, "synthesis_chars": 14000}


def _rank_relevant_literature(pdf_pages: pd.DataFrame | None, query: str, limit: int, snippet_chars: int) -> list[dict[str, Any]]:
    if pdf_pages is None or pdf_pages.empty or limit <= 0:
        return []
    ev = extract_source_evidence(pdf_pages)
    if ev.empty:
        return []
    corpus = (ev["document"].astype(str) + " p." + ev["page"].astype(str) + " " + ev["snippet"].fillna("").astype(str)).tolist()
    try:
        vec = TfidfVectorizer(ngram_range=(1, 2), strip_accents="unicode", stop_words="english")
        mat = vec.fit_transform(corpus + [str(query)])
        scores = cosine_similarity(mat[-1], mat[:-1]).ravel()
        order = np.argsort(scores)[::-1][: int(limit)]
    except Exception:
        order = np.arange(min(int(limit), len(ev)))
    out = []
    for idx in order:
        row = ev.iloc[int(idx)]
        snippet = re.sub(r"\s+", " ", str(row.get("snippet", "") or "")).strip()[: int(snippet_chars)]
        if snippet:
            out.append({"document": str(row.get("document", "")), "page": int(row.get("page", 0) or 0), "snippet": snippet})
    return out


def ai_research_question_prompt(
    df: pd.DataFrame,
    pdf_pages: pd.DataFrame | None,
    goal: str,
    count: int,
    existing: Sequence[str] = (),
    *,
    focus_columns: Sequence[str] = (),
    context_profile: str = "Compact",
) -> str:
    """Build a compact, evidence-grounded RQ prompt.

    Earlier versions serialised almost the whole 83+ variable schema and dozens
    of PDF passages.  That could exceed free-provider TPM limits before output
    generation began.  v5.8.5 keeps the selected research roles first and adds
    only a bounded amount of genuinely relevant context.
    """
    limits = _context_profile_limits(context_profile)
    numeric_all = list(df.select_dtypes(include=np.number).columns)
    focus = [str(c) for c in focus_columns if c is not None and str(c) in df.columns]
    focus = list(dict.fromkeys(focus))

    # Add a small number of semantically useful fields, never the full schema.
    semantic = [c for c in df.columns if any(tok in str(c).lower() for tok in ("year", "date", "time", "region", "nuts", "municip", "budget", "cost", "fund", "score", "criterion", "sector", "intervention"))]
    candidates = list(dict.fromkeys([*focus, *semantic, *numeric_all]))[: limits["schema"]]
    schema = [
        {"name": str(c), "dtype": str(df[c].dtype), "non_missing": int(df[c].notna().sum()), "unique": int(df[c].nunique(dropna=True))}
        for c in candidates
    ]

    numeric = [c for c in candidates if c in numeric_all]
    corr_summary: list[tuple[float, str, str, float]] = []
    if len(numeric) >= 2:
        try:
            corr = df[numeric].corr(numeric_only=True)
            for i, a in enumerate(corr.columns):
                for b in corr.columns[i + 1:]:
                    r = corr.loc[a, b]
                    if pd.notna(r):
                        corr_summary.append((abs(float(r)), str(a), str(b), float(r)))
            corr_summary = sorted(corr_summary, reverse=True)[: limits["corr"]]
        except Exception:
            corr_summary = []

    query = " ".join([str(goal)[: limits["goal_chars"]], *focus])
    literature = _rank_relevant_literature(pdf_pages, query, limits["literature"], limits["snippet_chars"])
    previous = [str(q)[:220] for q in list(existing)[-limits["existing"]:]]
    goal_text = re.sub(r"\s+", " ", str(goal)).strip()[: limits["goal_chars"]]

    return f"""Generate {int(count)} DISTINCT, specific, researchable questions for this project.
Ground every question in the compact evidence below. Use actual variable names where useful. Do not invent variables or source claims. Use descriptive, econometric, causal-design, Bayesian, predictive/XAI, spatial/time, MCDA/optimisation or robustness methods only when the supplied fields/evidence can support them.

GOAL: {goal_text}
FOCUS VARIABLES SELECTED BY USER: {json.dumps(focus, ensure_ascii=False)}
COMPACT SCHEMA: {json.dumps(schema, ensure_ascii=False)}
TOP OBSERVED CORRELATION LEADS (association only): {json.dumps(corr_summary, ensure_ascii=False)}
MOST RELEVANT PDF PASSAGES: {json.dumps(literature, ensure_ascii=False)}
RECENT QUESTIONS TO AVOID DUPLICATING: {json.dumps(previous, ensure_ascii=False)}

Return ONLY a JSON array. Each object must contain: research_question, method_family, variables, priority (1-3), rationale, source_basis."""


def estimate_rq_prompt_tokens(
    df: pd.DataFrame,
    pdf_pages: pd.DataFrame | None,
    goal: str,
    count: int,
    *,
    focus_columns: Sequence[str] = (),
    context_profile: str = "Compact",
) -> int:
    """Conservative visible estimate; no tokenizer dependency required."""
    prompt = ai_research_question_prompt(
        df, pdf_pages, goal, count, (), focus_columns=focus_columns, context_profile=context_profile
    )
    # English/JSON research prompts usually average below 4 chars/token; use 3.5
    # so the UI estimate is deliberately conservative.
    return int(math.ceil(len(prompt) / 3.5))

def generate_questions_with_ai(
    df: pd.DataFrame,
    pdf_pages: pd.DataFrame | None,
    goal: str,
    total: int,
    reply_fn: Callable[[str], str],
    batch_size: int = 25,
    *,
    focus_columns: Sequence[str] = (),
    context_profile: str = "Compact",
) -> pd.DataFrame:
    """Generate an AI-grounded RQ bank with resilient parsing and recovery.

    The function never discards a usable batch because a model ignored JSON
    formatting.  If an AI batch remains unusable after a constrained retry, the
    remaining slots are filled by the deterministic data-aware generator and the
    DataFrame attrs record how much came from each path.
    """
    total = min(max(1, int(total)), 200)
    rows: list[dict[str, Any]] = []
    seen: set[str] = set()
    attempts = 0
    parse_failures = 0
    max_attempts = max(3, math.ceil(total / max(1, batch_size)) * 3)

    while len(rows) < total and attempts < max_attempts:
        attempts += 1
        need = min(int(batch_size), total - len(rows))
        prompt = ai_research_question_prompt(df, pdf_pages, goal, need, [r["research_question"] for r in rows], focus_columns=focus_columns, context_profile=context_profile)
        response_text = reply_fn(prompt)
        parsed = _parse_rq_json(response_text)

        if not parsed:
            parse_failures += 1
            # One compact repair pass. Smaller batches reduce truncation risk and
            # make even non-schema providers much more reliable.
            repair_need = min(10, need)
            repair_prompt = ai_research_question_prompt(df, pdf_pages, goal, repair_need, [r["research_question"] for r in rows], focus_columns=focus_columns, context_profile=context_profile) + (
                "\n\nFORMAT RECOVERY: Your previous response could not be parsed. Return a raw JSON array only. "
                "Do not use Markdown fences, headings, commentary or trailing text."
            )
            parsed = _parse_rq_json(reply_fn(repair_prompt))

        for item in parsed:
            key = _normalise_query(item["research_question"])
            if key and key not in seen:
                seen.add(key)
                rows.append(item)
                if len(rows) >= total:
                    break

        # If two consecutive calls provide nothing useful, stop spending tokens
        # and recover deterministically rather than exposing an error to the user.
        if not parsed and parse_failures >= 2:
            break

    ai_count = len(rows)

    # Deterministic recovery preserves functionality under malformed output,
    # rate limits or weaker local models. It uses actual observed relationships,
    # not a generic placeholder question bank.
    if len(rows) < total:
        fallback = generate_research_questions(df, pdf_pages, limit=total)
        if isinstance(fallback, pd.DataFrame) and not fallback.empty:
            for item in fallback.to_dict("records"):
                q = str(item.get("research_question", "")).strip()
                key = _normalise_query(q)
                if not key or key in seen:
                    continue
                seen.add(key)
                rows.append({
                    "research_question": q,
                    "method_family": str(item.get("method_family", "Deterministic data-aware")),
                    "variables": str(item.get("variables", "")),
                    "priority": int(item.get("priority", 2)),
                    "rationale": str(item.get("rationale", "")),
                    "source_basis": str(item.get("source_basis", "Active dataset")),
                })
                if len(rows) >= total:
                    break

    if not rows:
        raise RuntimeError("Research-question generation produced no usable questions from either the selected AI engine or the deterministic recovery engine.")

    for i, row in enumerate(rows, 1):
        row["rq_id"] = f"RQ{i:03d}"
    cols = ["rq_id", "research_question", "method_family", "variables", "priority", "rationale", "source_basis"]
    frame = pd.DataFrame(rows)[cols]
    frame.attrs["ai_generated"] = int(ai_count)
    frame.attrs["deterministic_recovery"] = int(len(frame) - ai_count)
    frame.attrs["parse_failures"] = int(parse_failures)
    return frame

def _offline_narrative(tables: dict[str, pd.DataFrame], plan: AgenticPlan) -> dict[str, str]:
    desc = tables.get("Descriptive statistics", pd.DataFrame())
    top = tables.get("Top correlations", pd.DataFrame())
    coef = tables.get("OLS coefficients", pd.DataFrame())
    quality = tables.get("Quality audit", pd.DataFrame())
    paragraphs = []
    if not quality.empty:
        pairs = "; ".join(f"{r.check}: {r.value}" for r in quality.head(6).itertuples(index=False))
        paragraphs.append(f"The automated audit characterised the active analytical scope before modelling ({pairs}).")
    lit = tables.get("Literature source evidence", pd.DataFrame())
    terms = tables.get("Literature key terms", pd.DataFrame())
    if not lit.empty:
        docs = int(lit.document.nunique()) if "document" in lit else 0
        pages = len(lit)
        topic_text = ", ".join(terms.term.head(8).astype(str)) if not terms.empty else "no stable key terms extracted"
        paragraphs.append(f"The local literature pass indexed {docs} document(s) across {pages} extracted pages. Frequent source terms included {topic_text}. These are extractive navigation signals rather than verified claims or complete bibliographic records.")
    if not desc.empty:
        most_variable = desc.assign(abs_cv=pd.to_numeric(desc.cv, errors="coerce").abs()).sort_values("abs_cv", ascending=False).head(1)
        if not most_variable.empty and pd.notna(most_variable.iloc[0].abs_cv):
            r = most_variable.iloc[0]
            paragraphs.append(f"Among the screened numeric variables, {r.variable} showed the greatest relative dispersion (|CV|={r.abs_cv:.3f}); distributional shape and outliers should therefore be reported alongside the mean.")
    if not top.empty:
        r = top.iloc[0]
        paragraphs.append(f"The strongest screened pairwise association was between {r.variable_1} and {r.variable_2} (r={r.correlation:.3f}, p={r.p_value:.4g}). This is an association screen, not a causal estimate.")
    if not coef.empty:
        nonconst = coef[~coef.term.astype(str).str.lower().isin(["const", "intercept"])]
        if not nonconst.empty:
            r = nonconst.iloc[np.nanargmin(pd.to_numeric(nonconst.p_value, errors="coerce").fillna(1).to_numpy())]
            paragraphs.append(f"In the configured primary OLS model, the lowest p-value among non-intercept terms corresponded to {r.term} (β={r.coefficient:.4g}, 95% CI {r.ci_95_low:.4g} to {r.ci_95_high:.4g}, p={r.p_value:.4g}). Interpretation remains conditional on model specification and diagnostics.")
    if not paragraphs:
        paragraphs.append("The deterministic workflow completed the available audit and evidence-indexing steps. Additional model interpretation requires explicit variable mappings or a richer analytical scope.")
    discussion = "\n\n".join(paragraphs)
    conclusion = (f"For the research goal — {plan.goal or 'the configured research objective'} — the automated workflow provides a reproducible first-pass empirical synthesis. "
                  "The strongest defensible conclusions are those directly supported by the exported tables and diagnostics. Causal language, exact bibliographic claims and final submission formatting require explicit verification.")
    limitations = "The automated package is designed to remove repetitive analytical work, not scholarly responsibility. Missing data, non-random sampling, measurement validity, model specification, multiple testing, external validity and causal identification must be reviewed before submission."
    return {"discussion": discussion, "conclusion": conclusion, "limitations": limitations}


def run_agentic_workflow(
    df: pd.DataFrame,
    *,
    plan: AgenticPlan,
    pdf_pages: pd.DataFrame | None = None,
    outcome: str | None = None,
    predictors: Sequence[str] = (),
    group: str | None = None,
    question_limit: int = 150,
) -> AgenticRun:
    tables: dict[str, pd.DataFrame] = {}
    tables["Quality audit"] = quality_summary(df)
    numeric = list(df.select_dtypes(include=np.number).columns)
    selected_numeric = list(dict.fromkeys([*(predictors or []), *([outcome] if outcome else [])]))
    selected_numeric = [c for c in selected_numeric if c in numeric]
    if len(selected_numeric) < 2:
        selected_numeric = numeric[:min(30, len(numeric))]
    tables["Descriptive statistics"] = descriptive_statistics(df, selected_numeric or numeric[:30])
    if selected_numeric:
        corr, pvals = correlation_matrix(df, selected_numeric, method="pearson")
        tables["Correlation matrix"] = corr.reset_index(names="variable")
        tables["Correlation p-values"] = pvals.reset_index(names="variable")
        tables["Top correlations"] = _top_correlations(corr, pvals)
        static, interactive = _correlation_figure_bytes(corr)
    else:
        corr = pd.DataFrame(); static, interactive = {}, {}
    ols_static, ols_interactive = {}, {}
    if outcome and predictors:
        try:
            model = fit_detailed_model(df, outcome, list(predictors), estimator="OLS", covariance="HC3")
            tables["OLS coefficients"] = model.coefficients
            tables["OLS fit"] = model.fit
            tables["OLS diagnostics"] = model.diagnostics
            tables["OLS predictions"] = model.predictions
            ols_static, ols_interactive = _ols_figure_bytes(model.coefficients)
        except Exception as exc:
            ols_static, ols_interactive = {}, {}
            tables["OLS execution note"] = pd.DataFrame([{"status": "not executed", "reason": str(exc)}])
    if group and outcome:
        try:
            tables["Group tests"] = group_tests(df, [outcome], group)
            tables["Group frequencies"] = categorical_summary(df, [group])
        except Exception as exc:
            tables["Group execution note"] = pd.DataFrame([{"status": "not executed", "reason": str(exc)}])
    if len(selected_numeric) >= 2 and len(df) >= 10:
        try:
            loadings, variance = pca_table(df, selected_numeric[:20], n_components=min(5, len(selected_numeric)))
            tables["PCA loadings"] = loadings
            tables["PCA variance"] = variance
        except Exception as exc:
            tables["PCA execution note"] = pd.DataFrame([{"status": "not executed", "reason": str(exc)}])
        try:
            assignments, profiles = cluster_table(df, selected_numeric[:12], clusters=min(4, max(2, len(df)//20)))
            tables["Cluster assignments"] = assignments
            tables["Cluster profiles"] = profiles
        except Exception as exc:
            tables["Cluster execution note"] = pd.DataFrame([{"status": "not executed", "reason": str(exc)}])
    source_evidence = extract_source_evidence(pdf_pages if pdf_pages is not None else pd.DataFrame())
    tables["Literature source evidence"] = source_evidence
    tables["Literature key terms"] = literature_key_terms(pdf_pages if pdf_pages is not None else pd.DataFrame())
    tables["Research questions"] = generate_research_questions(df, pdf_pages, limit=question_limit)
    static.update(ols_static)
    interactive.update(ols_interactive)
    narratives = _offline_narrative(tables, plan)
    manifest = {
        "version": "5.8.2",
        "mode": "offline deterministic agentic workflow",
        "goal": plan.goal,
        "mappings": plan.mappings,
        "steps": plan.steps,
        "warnings": plan.warnings,
        "rows": int(len(df)),
        "columns": int(df.shape[1]),
        "tables_generated": sorted(tables),
        "research_questions_generated": int(len(tables["Research questions"])),
        "pdf_documents": sorted(pdf_pages.document.unique().tolist()) if pdf_pages is not None and not pdf_pages.empty else [],
    }
    return AgenticRun(plan=plan, tables=tables, narratives=narratives, figures=static, interactive_html=interactive, manifest=manifest)


def _paper_markdown(run: AgenticRun, title: str = "Agentic Research Draft") -> str:
    q = run.tables.get("Research questions", pd.DataFrame()).head(10)
    q_text = "\n".join(f"- {r.research_question}" for r in q.itertuples(index=False)) or "- No research questions generated."
    source = run.tables.get("Literature source evidence", pd.DataFrame())
    source_lines = []
    if not source.empty:
        for row in source.head(60).itertuples(index=False):
            doi = f" DOI: {row.doi}." if getattr(row, "doi", "") else ""
            source_lines.append(f"- {row.document}, p. {row.page}.{doi} {row.snippet[:300]}")
    sources = "\n".join(source_lines) or "- No PDF evidence supplied."
    return f"""# {title}

## Status

Near-submission analytical draft generated by Makryvelios Agentic Research Mode. Human verification of claims, references, formatting and institutional requirements remains mandatory.

## Abstract

{run.narratives.get('abstract', '[AI synthesis not run; see deterministic outputs and discussion below.]')}

## Research objective

{run.plan.goal or '[Define the principal research objective]'}

## Candidate research questions

{q_text}

## Data and methods

The workflow audited the active dataset, produced descriptive and association evidence, and executed only analytical routines supported by the configured variable mappings. Where a primary outcome and predictors were supplied, HC3-robust OLS was estimated; PCA and clustering were run where the numeric matrix was sufficient. All tables, diagnostics, mappings and workflow steps are included in the reproducibility package.

## Results

{run.narratives.get('results', run.narratives.get('discussion',''))}

## Discussion

The reported patterns should be interpreted in relation to the substantive research design and uploaded literature. Automated synthesis prioritises reproducible numerical evidence and explicitly separates association, prediction, optimisation and causal claims.

## Conclusion

{run.narratives.get('conclusion','')}

## Limitations

{run.narratives.get('limitations','')}

## Literature evidence notes

{sources}

## Submission checklist

1. Verify every bibliographic reference and quotation against the original PDF.
2. Confirm the final research question, hypotheses and causal identification language.
3. Review model diagnostics, missing-data handling and multiplicity.
4. Apply the institution/journal style guide and final reference manager output.
5. Re-run the workflow after any data or specification change and retain the exported manifest.
"""


def _agentic_docx_bytes(run: AgenticRun, title: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("Word export requires python-docx from the bundled requirements.") from exc
    document = Document()
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"
            if style_name == "Normal":
                styles[style_name].font.size = Pt(11)
    document.add_heading(title, 0)
    p = document.add_paragraph()
    r = p.add_run("Near-submission analytical draft generated by Makryvelios Agentic Research Mode. Human verification is required before submission.")
    r.bold = True
    document.add_heading("Abstract", level=1)
    document.add_paragraph(run.narratives.get("abstract", "AI synthesis was not run; the package therefore retains the deterministic first-pass narrative."))
    document.add_heading("Research objective", level=1)
    document.add_paragraph(run.plan.goal or "[Define the principal research objective]")
    document.add_heading("Candidate research questions", level=1)
    rqs = run.tables.get("Research questions", pd.DataFrame()).head(15)
    if rqs.empty:
        document.add_paragraph("No research questions were generated.")
    else:
        for row in rqs.itertuples(index=False):
            document.add_paragraph(f"{row.rq_id}. {row.research_question}", style="List Number")
    document.add_heading("Data and methods", level=1)
    document.add_paragraph("The approved workflow audited the active dataset, indexed selected literature evidence locally, generated a candidate research-question bank and executed only analytical routines supported by the mapped variables. Numerical outputs are deterministic and do not require an external LLM API.")
    plan_table = document.add_table(rows=1, cols=4)
    plan_table.style = "Table Grid"
    for i, h in enumerate(["Step", "Action", "Engine", "Output"]): plan_table.rows[0].cells[i].text = h
    for item in run.plan.steps:
        cells = plan_table.add_row().cells
        for i, key in enumerate(["step", "action", "engine", "output"]): cells[i].text = str(item.get(key, ""))
    document.add_heading("Results", level=1)
    document.add_paragraph(run.narratives.get("results", run.narratives.get("discussion", "")))
    for table_name in ["OLS coefficients", "OLS fit", "OLS diagnostics", "Top correlations", "Group tests", "PCA variance"]:
        table = run.tables.get(table_name, pd.DataFrame())
        if table is None or table.empty:
            continue
        document.add_heading(table_name, level=2)
        clipped = table.head(25).copy()
        doc_table = document.add_table(rows=1, cols=len(clipped.columns))
        doc_table.style = "Table Grid"
        for j, col in enumerate(clipped.columns): doc_table.rows[0].cells[j].text = str(col)
        for _, row in clipped.iterrows():
            cells = doc_table.add_row().cells
            for j, col in enumerate(clipped.columns):
                value = row[col]
                cells[j].text = "" if pd.isna(value) else (f"{value:.5g}" if isinstance(value, (float, np.floating)) else str(value))
    png = run.figures.get("correlation_matrix.png") or run.figures.get("ols_coefficients.png")
    if png:
        document.add_heading("Selected analytical figure", level=2)
        document.add_picture(io.BytesIO(png), width=Inches(6.3))
    document.add_heading("Discussion", level=1)
    document.add_paragraph(run.narratives.get("discussion", ""))
    document.add_heading("Conclusion", level=1)
    document.add_paragraph(run.narratives.get("conclusion", ""))
    document.add_heading("Limitations", level=1)
    document.add_paragraph(run.narratives.get("limitations", ""))
    document.add_heading("Literature evidence notes", level=1)
    evidence = run.tables.get("Literature source evidence", pd.DataFrame()).head(40)
    if evidence.empty:
        document.add_paragraph("No PDF evidence was supplied.")
    else:
        for row in evidence.itertuples(index=False):
            doi = f" DOI: {row.doi}." if getattr(row, "doi", "") else ""
            document.add_paragraph(f"{row.document}, p. {row.page}.{doi} {str(row.snippet)[:450]}", style="List Bullet")
    document.add_heading("Submission verification checklist", level=1)
    for item in [
        "Verify every bibliographic reference and quotation against the original PDF.",
        "Confirm the final research question, hypotheses and any causal-identification language.",
        "Review model diagnostics, missing-data handling, multiplicity and robustness.",
        "Apply the institution/journal style guide and final reference-manager output.",
        "Re-run after any data/specification change and retain the matching manifest.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    out = io.BytesIO(); document.save(out); return out.getvalue()


def _html_report(run: AgenticRun, title: str) -> bytes:
    sections = [f"<h1>{html.escape(title)}</h1>", "<p><b>Near-submission analytical draft; human verification required.</b></p>"]
    sections.append("<h2>Abstract</h2><p>" + html.escape(run.narratives.get("abstract", "AI synthesis not run.")).replace("\n", "<br>") + "</p>")
    sections.append(f"<h2>Research objective</h2><p>{html.escape(run.plan.goal)}</p>")
    sections.append("<h2>Results</h2><p>" + html.escape(run.narratives.get("results", run.narratives.get("discussion", ""))).replace("\n", "<br>") + "</p>")
    sections.append("<h2>Discussion</h2><p>" + html.escape(run.narratives.get("discussion", "")).replace("\n", "<br>") + "</p>")
    sections.append("<h2>Conclusion</h2><p>" + html.escape(run.narratives.get("conclusion", "")).replace("\n", "<br>") + "</p>")
    for name, table in run.tables.items():
        if table is None or table.empty:
            continue
        sections.append(f"<h2>{html.escape(name)}</h2>")
        sections.append(table.head(500).to_html(index=False, escape=True))
    for name, payload in run.interactive_html.items():
        # Include linkable self-contained interactive files in ZIP rather than nesting huge JS here.
        sections.append(f"<p>Interactive figure included in package: <code>interactive/{html.escape(name)}</code></p>")
    css = "body{font-family:Georgia,serif;max-width:1300px;margin:40px auto;padding:0 24px;color:#172033}table{border-collapse:collapse;width:100%;font-size:12px}th,td{border:1px solid #ccd5df;padding:6px;vertical-align:top}th{background:#e9eef4}h1,h2{color:#173a63}code{background:#f1f5f9;padding:2px 4px}"
    return ("<!doctype html><html><head><meta charset='utf-8'><title>" + html.escape(title) + "</title><style>" + css + "</style></head><body>" + "\n".join(sections) + "</body></html>").encode("utf-8")


def agentic_submission_package(run: AgenticRun, title: str = "Agentic Research Draft") -> bytes:
    """Build a complete near-submission package: DOCX, XLSX, JSON, HTML and figures."""
    paper = _paper_markdown(run, title=title)
    workbook = to_excel_bytes({name[:31]: table for name, table in run.tables.items() if isinstance(table, pd.DataFrame)})
    report = _html_report(run, title)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("paper/paper_draft.docx", _agentic_docx_bytes(run, title))
        archive.writestr("paper/paper_draft.md", paper.encode("utf-8"))
        archive.writestr("results/complete_results.xlsx", workbook)
        archive.writestr("results/manifest.json", json.dumps(run.manifest, ensure_ascii=False, indent=2).encode("utf-8"))
        archive.writestr("report/interactive_report.html", report)
        for name, payload in run.figures.items():
            archive.writestr(f"figures/{name}", payload)
        for name, payload in run.interactive_html.items():
            archive.writestr(f"interactive/{name}", payload)
        archive.writestr("README.txt", (
            "Makryvelios Agentic Research Mode v5.8.2\n\n"
            "This package is generated locally from the configured data and PDF evidence.\n"
            "Numerical outputs do not require an external AI API.\n"
            "The draft is designed to be close to submission-ready but must be checked by a human researcher before submission.\n"
            "Verify references, quotations, causal claims, institutional formatting and every substantive conclusion.\n"
        ).encode("utf-8"))
    return buf.getvalue()
