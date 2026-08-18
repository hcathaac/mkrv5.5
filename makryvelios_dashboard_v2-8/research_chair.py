"""Free/offline research-command engine for the Makryvelios dashboard.

The module is deliberately independent of paid APIs.  It provides deterministic
data/PDF scoping, safe formula evaluation, reproducible protocol execution,
natural-language summaries and paper-blueprint exports.  A local Ollama server
may optionally improve the prose, but is never required.
"""
from __future__ import annotations

import ast
import html
import io
import json
import math
import re
import zipfile
from dataclasses import dataclass, field
from typing import Any

import numpy as np
import pandas as pd
from scipy import stats

from output_guidance import build_output_guide, output_guide_markdown
from prompt_library import prompt_library, prompt_library_markdown


@dataclass
class ProtocolResult:
    algorithm: str
    tables: dict[str, pd.DataFrame] = field(default_factory=dict)
    comments: list[str] = field(default_factory=list)
    equation: str = ""
    executed_expression: str = ""


@dataclass
class FeasibilityDecision:
    """Transparent pre-execution verdict for a natural-language command."""

    status: str
    reason: str
    route: str
    can_execute: bool


def extract_pdf_pages(name: str, payload: bytes) -> list[dict[str, Any]]:
    """Extract page-level text from a PDF without sending it off-device."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency message is UI-facing
        raise RuntimeError("PDF support requires pypdf. Install the bundled requirements.") from exc
    reader = PdfReader(io.BytesIO(payload))
    pages: list[dict[str, Any]] = []
    for number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        pages.append({"document": name, "page": number, "text": text, "characters": len(text)})
    return pages


def extract_pdf_collection(items: tuple[tuple[str, bytes], ...]) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for name, payload in items:
        rows.extend(extract_pdf_pages(name, payload))
    return pd.DataFrame(rows, columns=["document", "page", "text", "characters"])


def select_pdf_evidence(
    pages: pd.DataFrame,
    documents: list[str] | None = None,
    page_ranges: dict[str, tuple[int, int]] | None = None,
    keywords: str = "",
    max_characters: int = 60_000,
) -> pd.DataFrame:
    """Return only explicitly selected PDF documents/pages/keyword passages."""
    if pages.empty:
        return pages.copy()
    out = pages.copy()
    if documents is not None:
        out = out[out["document"].isin(documents)]
    if page_ranges:
        keep = pd.Series(False, index=out.index)
        for document, (start, end) in page_ranges.items():
            keep |= (out.document == document) & out.page.between(int(start), int(end))
        out = out[keep]
    terms = [term.strip() for term in re.split(r"[,;\n]+", keywords) if term.strip()]
    if terms:
        pattern = "|".join(re.escape(term) for term in terms)
        out = out[out.text.str.contains(pattern, case=False, regex=True, na=False)]
    if max_characters > 0 and not out.empty:
        eligible = out.copy()
        cumulative = out.characters.fillna(0).astype(int).cumsum()
        out = out[cumulative <= max_characters]
        if out.empty and not eligible.empty:
            out = eligible.head(1).copy()
            out["text"] = out["text"].str.slice(0, max_characters)
            out["characters"] = out["text"].str.len()
    return out.reset_index(drop=True)


def _coerce_year(series: pd.Series) -> pd.Series:
    numeric = pd.to_numeric(series, errors="coerce")
    plausible = numeric.between(1800, 2200)
    if plausible.sum() >= max(2, int(series.notna().sum() * .5)):
        return numeric.where(plausible)
    parsed = pd.to_datetime(series, errors="coerce")
    return parsed.dt.year.astype(float)


def year_bounds(df: pd.DataFrame, year_column: str | None) -> tuple[int, int] | None:
    if not year_column or year_column not in df:
        return None
    years = _coerce_year(df[year_column]).dropna()
    if years.empty:
        return None
    return int(years.min()), int(years.max())


def apply_scope(
    df: pd.DataFrame,
    selected_columns: list[str] | None = None,
    year_column: str | None = None,
    start_year: int | None = None,
    end_year: int | None = None,
    categorical_filters: dict[str, list[Any]] | None = None,
    numeric_filters: dict[str, tuple[float, float]] | None = None,
) -> pd.DataFrame:
    out = df.copy()
    if year_column and year_column in out and start_year is not None and end_year is not None:
        years = _coerce_year(out[year_column])
        out = out[years.between(start_year, end_year)]
    for column, values in (categorical_filters or {}).items():
        if column in out and values:
            out = out[out[column].isin(values)]
    for column, bounds in (numeric_filters or {}).items():
        if column in out:
            values = pd.to_numeric(out[column], errors="coerce")
            out = out[values.between(float(bounds[0]), float(bounds[1]))]
    if selected_columns:
        ordered = list(dict.fromkeys([c for c in selected_columns if c in out]))
        if year_column and year_column in out and year_column not in ordered:
            ordered.insert(0, year_column)
        out = out[ordered]
    return out.reset_index(drop=True)


_BINARY = {
    ast.Add: lambda a, b: a + b,
    ast.Sub: lambda a, b: a - b,
    ast.Mult: lambda a, b: a * b,
    ast.Div: lambda a, b: a / b,
    ast.Pow: lambda a, b: a**b,
    ast.Mod: lambda a, b: a % b,
}
_UNARY = {ast.UAdd: lambda a: a, ast.USub: lambda a: -a}
_FUNCTIONS = {
    "log": np.log,
    "log1p": np.log1p,
    "exp": np.exp,
    "sqrt": np.sqrt,
    "abs": np.abs,
}


def _eval_node(node: ast.AST, frame: pd.DataFrame):
    if isinstance(node, ast.Expression):
        return _eval_node(node.body, frame)
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return float(node.value)
    if isinstance(node, ast.Name):
        if node.id not in frame:
            raise ValueError(f"Unknown variable in expression: {node.id}")
        return pd.to_numeric(frame[node.id], errors="coerce")
    if isinstance(node, ast.BinOp) and type(node.op) in _BINARY:
        return _BINARY[type(node.op)](_eval_node(node.left, frame), _eval_node(node.right, frame))
    if isinstance(node, ast.UnaryOp) and type(node.op) in _UNARY:
        return _UNARY[type(node.op)](_eval_node(node.operand, frame))
    if isinstance(node, ast.Call) and isinstance(node.func, ast.Name) and node.func.id in _FUNCTIONS and len(node.args) == 1:
        return _FUNCTIONS[node.func.id](_eval_node(node.args[0], frame))
    raise ValueError("Only numeric columns, constants, + - * / ** %, parentheses, log, log1p, exp, sqrt and abs are permitted.")


def add_safe_derived_column(df: pd.DataFrame, name: str, expression: str) -> pd.DataFrame:
    """Evaluate a restricted mathematical expression; never execute arbitrary code."""
    if not name or not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", name):
        raise ValueError("Derived-variable name must use letters, numbers and underscores and cannot begin with a number.")
    if len(expression) > 1_000:
        raise ValueError("Expression is too long.")
    tree = ast.parse(expression, mode="eval")
    out = df.copy()
    values = _eval_node(tree, out)
    if np.isscalar(values):
        values = pd.Series(values, index=out.index)
    out[name] = pd.to_numeric(values, errors="coerce").replace([np.inf, -np.inf], np.nan)
    return out


def scope_profile(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame([{"records": 0, "variables": 0, "numeric_variables": 0, "missing_cells": 0, "missing_percent": np.nan}])
    missing = int(df.isna().sum().sum())
    cells = max(int(df.shape[0] * df.shape[1]), 1)
    return pd.DataFrame([{
        "records": len(df),
        "variables": df.shape[1],
        "numeric_variables": len(df.select_dtypes(include=np.number).columns),
        "missing_cells": missing,
        "missing_percent": 100 * missing / cells,
        "duplicate_rows": int(df.duplicated().sum()),
    }])


def _descriptive_table(df: pd.DataFrame, variables: list[str]) -> pd.DataFrame:
    chosen = [c for c in variables if c in df and pd.api.types.is_numeric_dtype(df[c])]
    if not chosen:
        return pd.DataFrame()
    table = df[chosen].describe(percentiles=[.25, .5, .75]).T.reset_index(names="variable")
    table["missing"] = [int(df[c].isna().sum()) for c in chosen]
    table["missing_percent"] = [100 * float(df[c].isna().mean()) for c in chosen]
    return table


def _missingness_table(df: pd.DataFrame) -> pd.DataFrame:
    """Return a variable-level audit even when no model variables were chosen."""
    if df.shape[1] == 0:
        return pd.DataFrame()
    rows = []
    for column in df.columns:
        missing = int(df[column].isna().sum())
        rows.append({
            "variable": column,
            "dtype": str(df[column].dtype),
            "observed": int(df[column].notna().sum()),
            "missing": missing,
            "missing_percent": 100 * missing / max(len(df), 1),
            "unique_observed": int(df[column].nunique(dropna=True)),
        })
    return pd.DataFrame(rows).sort_values(
        ["missing_percent", "variable"], ascending=[False, True]
    ).reset_index(drop=True)


def _correlation_table(df: pd.DataFrame, variables: list[str], maximum: int = 25) -> tuple[pd.DataFrame, int]:
    """Bound pairwise screening so a 1,000-variable scope remains responsive."""
    chosen = [c for c in variables if c in df and pd.api.types.is_numeric_dtype(df[c])]
    total = len(chosen)
    chosen = chosen[:maximum]
    rows: list[dict[str, Any]] = []
    for position, left in enumerate(chosen):
        for right in chosen[position + 1:]:
            pair = df[[left, right]].dropna()
            if len(pair) < 3 or pair[left].nunique() < 2 or pair[right].nunique() < 2:
                continue
            coefficient, p_value = stats.spearmanr(pair[left], pair[right])
            rows.append({
                "variable_1": left,
                "variable_2": right,
                "n": len(pair),
                "spearman_rho": coefficient,
                "p_value": p_value,
            })
    table = pd.DataFrame(rows)
    if not table.empty:
        table["absolute_rho"] = table["spearman_rho"].abs()
        table = table.sort_values(["absolute_rho", "p_value"], ascending=[False, True]).reset_index(drop=True)
    return table, total


def _longitudinal_table(
    df: pd.DataFrame,
    variables: list[str],
    year_column: str,
    group_column: str | None,
    aggregation: str,
) -> pd.DataFrame:
    numeric = [
        c for c in variables
        if c in df and c != year_column and pd.api.types.is_numeric_dtype(df[c])
    ]
    if not numeric:
        return pd.DataFrame()
    work = df.copy()
    work["__analysis_year__"] = _coerce_year(work[year_column])
    keys = ["__analysis_year__"] + ([group_column] if group_column and group_column in work else [])
    agg_map = {"Mean": "mean", "Sum": "sum", "Median": "median", "Count": "count"}
    trend = (
        work.dropna(subset=["__analysis_year__"])
        .groupby(keys, dropna=False)[numeric]
        .agg(agg_map.get(aggregation, "mean"))
        .reset_index()
        .rename(columns={"__analysis_year__": "year"})
    )
    return trend.sort_values(["year"] + ([group_column] if group_column and group_column in trend else [])).reset_index(drop=True)


def execute_protocol(
    df: pd.DataFrame,
    algorithm: str,
    outcome: str | None = None,
    predictors: list[str] | None = None,
    year_column: str | None = None,
    group_column: str | None = None,
    aggregation: str = "Mean",
    equation: str = "",
    executed_expression: str = "",
) -> ProtocolResult:
    predictors = [c for c in (predictors or []) if c in df]
    result = ProtocolResult(algorithm=algorithm, equation=equation, executed_expression=executed_expression)
    result.tables["Scope profile"] = scope_profile(df)
    variables = list(dict.fromkeys(([outcome] if outcome in df else []) + predictors))
    # A descriptive command without an explicit model specification means
    # "describe the selected scope", not "return an empty table".
    if not variables:
        variables = list(df.select_dtypes(include=np.number).columns)
    numeric = [c for c in variables if pd.api.types.is_numeric_dtype(df[c])]
    result.tables["Descriptive statistics"] = _descriptive_table(df, numeric)
    result.tables["Variable missingness"] = _missingness_table(df)

    if algorithm == "Longitudinal trend":
        if not year_column or year_column not in df:
            raise ValueError("Choose a valid year column for longitudinal analysis.")
        if not numeric:
            raise ValueError("Choose at least one numeric outcome or predictor.")
        trend = _longitudinal_table(df, numeric, year_column, group_column, aggregation)
        result.tables["Longitudinal results"] = trend
        result.comments.append(f"The table follows {len(trend.year.unique()) if not trend.empty else 0} observed years using {aggregation.lower()} aggregation.")
        result.comments.append("Temporal movement is descriptive unless the protocol supplies a credible identification strategy and suitable controls.")

    elif algorithm == "Correlation screening":
        if len(numeric) < 2:
            raise ValueError("Correlation screening requires at least two numeric variables.")
        correlations, total = _correlation_table(df, numeric)
        result.tables["Correlation screening"] = correlations
        result.comments.append("Spearman correlation detects monotonic association; it does not adjust for confounding or establish direction of effect.")
        if total > 25:
            result.comments.append(f"Responsiveness safeguard: pairwise screening used the first 25 of {total} selected numeric variables. Narrow the variable scope to test a different set.")

    elif algorithm == "OLS specification":
        if not outcome or outcome not in df or not predictors:
            raise ValueError("OLS requires one outcome and at least one predictor.")
        model_data = df[[outcome] + predictors].apply(pd.to_numeric, errors="coerce").dropna()
        if len(model_data) <= len(predictors) + 2:
            raise ValueError("Insufficient complete observations for the requested OLS specification.")
        y_values = model_data[outcome].to_numpy(float)
        x_values = np.column_stack([np.ones(len(model_data)), model_data[predictors].to_numpy(float)])
        terms = ["const"] + predictors
        bread = np.linalg.pinv(x_values.T @ x_values)
        beta = bread @ x_values.T @ y_values
        fitted = x_values @ beta
        residual = y_values - fitted
        leverage = np.einsum("ij,jk,ik->i", x_values, bread, x_values)
        adjusted_residual = residual / np.clip(1 - leverage, 1e-10, None)
        meat = x_values.T @ ((adjusted_residual**2)[:, None] * x_values)
        covariance = bread @ meat @ bread
        robust_se = np.sqrt(np.clip(np.diag(covariance), 0, None))
        t_values = np.divide(beta, robust_se, out=np.full_like(beta, np.nan), where=robust_se > 0)
        degrees_freedom = len(model_data) - x_values.shape[1]
        p_values = 2 * stats.t.sf(np.abs(t_values), df=max(degrees_freedom, 1))
        critical = stats.t.ppf(.975, df=max(degrees_freedom, 1))
        ci_low, ci_high = beta - critical * robust_se, beta + critical * robust_se
        rss = float(residual @ residual)
        centred = y_values - y_values.mean()
        tss = float(centred @ centred)
        r_squared = 1 - rss / tss if tss > 0 else np.nan
        adjusted_r_squared = 1 - (1 - r_squared) * (len(model_data) - 1) / max(degrees_freedom, 1) if np.isfinite(r_squared) else np.nan
        sigma2 = max(rss / len(model_data), np.finfo(float).tiny)
        log_likelihood = -.5 * len(model_data) * (math.log(2 * math.pi) + 1 + math.log(sigma2))
        parameter_count = x_values.shape[1]
        coefficients = pd.DataFrame({
            "term": terms,
            "coefficient": beta,
            "robust_se": robust_se,
            "t_value": t_values,
            "p_value": p_values,
            "ci_95_low": ci_low,
            "ci_95_high": ci_high,
        })
        fit_table = pd.DataFrame([{
            "n": len(model_data), "r_squared": r_squared,
            "adjusted_r_squared": adjusted_r_squared,
            "aic": -2 * log_likelihood + 2 * parameter_count,
            "bic": -2 * log_likelihood + math.log(len(model_data)) * parameter_count,
            "covariance": "HC3",
        }])
        result.tables["OLS coefficients"] = coefficients
        result.tables["OLS fit"] = fit_table
        result.comments.append("Coefficients are conditional associations under the stated specification; HC3 protects inference against general heteroskedasticity.")
        result.comments.append("Causal language requires defensible temporal ordering, measurement and identification beyond model fit.")

    elif algorithm == "Descriptive profile":
        result.comments.append("The output characterises the selected analytical sample and should precede confirmatory modelling.")

    else:
        result.comments.append("The custom algorithm has been documented but not interpreted as executable computer code. Only the safe data audit and requested derived expression were run.")
        result.comments.append("Implement custom statistical steps in a validated analytical module before treating them as estimated evidence.")
    return result


def _compact_table(table: pd.DataFrame, rows: int = 12) -> str:
    if table is None or table.empty:
        return "No rows available."
    return table.head(rows).to_csv(index=False)


def evidence_text(evidence: pd.DataFrame, max_chars: int = 16_000) -> str:
    if evidence.empty:
        return ""
    parts = []
    for row in evidence.itertuples(index=False):
        parts.append(f"[{row.document}, p. {row.page}]\n{row.text}")
    return "\n\n".join(parts)[:max_chars]


def build_offline_reply(
    question: str,
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
) -> str:
    """Produce a transparent natural-language answer without a generative API."""
    q = question.lower().strip()
    profile = result.tables.get("Scope profile", pd.DataFrame())
    records = int(profile.iloc[0].records) if not profile.empty else 0
    variables = int(profile.iloc[0].variables) if not profile.empty else 0
    opening = f"The answer is based on the explicitly selected analytical scope: {records:,} records and {variables:,} variables."
    if any(term in q for term in ["trend", "year", "χρον", "έτος", "ετ"]):
        trend = result.tables.get("Longitudinal results", pd.DataFrame())
        if trend.empty:
            body = "No longitudinal table has yet been generated. Select a year variable and run the Longitudinal trend protocol first."
        else:
            years = pd.to_numeric(trend.year, errors="coerce").dropna()
            body = f"The selected series covers {int(years.min())}–{int(years.max())} across {len(years.unique())} observed years. Inspect the downloadable longitudinal table for magnitude, breaks and missing years; movement alone is not a causal effect."
    elif any(term in q for term in ["coefficient", "regression", "ols", "significant", "παλινδ", "συντελεστ"]):
        coef = result.tables.get("OLS coefficients", pd.DataFrame())
        if coef.empty:
            body = "No OLS result is currently in the research context. Run the OLS specification protocol with an outcome and predictors."
        else:
            terms = coef[coef.term != "const"].sort_values("p_value").head(3)
            statements = [f"{r.term}: b={r.coefficient:.4g}, 95% CI [{r.ci_95_low:.4g}, {r.ci_95_high:.4g}], p={r.p_value:.4g}" for r in terms.itertuples()]
            body = "The most precisely estimated terms are: " + "; ".join(statements) + ". These are conditional associations, not automatic causal effects."
    elif any(term in q for term in ["missing", "quality", "clean", "ελλιπ", "ποιότη"]):
        if profile.empty:
            body = "No scoped data profile is available."
        else:
            row = profile.iloc[0]
            body = f"The selected scope contains {int(row.missing_cells):,} missing cells ({float(row.missing_percent):.2f}%) and {int(row.duplicate_rows):,} duplicate rows. Review whether missingness is structural, random or created by joins before modelling."
    elif any(term in q for term in ["pdf", "note", "literature", "θεωρ", "βιβλιο"]):
        if evidence.empty:
            body = "No PDF passages are selected. Upload PDFs and select documents, page ranges or keywords before asking an evidence-grounded question."
        else:
            refs = ", ".join(f"{r.document} p.{r.page}" for r in evidence.head(5).itertuples())
            body = f"The active documentary evidence contains {len(evidence)} selected pages. The first relevant locations are {refs}. Treat these as source notes and verify quotations against the original pages."
    elif any(term in q for term in ["paper", "article", "publish", "γράψ", "δημοσί"]):
        body = "Frame one principal research question, pre-state the outcome and model, report the analytical sample and exclusions, present estimates with uncertainty and diagnostics, separate association from causality, and place robustness checks before the conclusion. The downloadable paper blueprint converts the current protocol into a section-by-section manuscript plan."
    else:
        body = f"The active protocol is “{result.algorithm}”. Its equation is recorded as {result.equation or 'not specified'}. The generated tables and PDF evidence delimit what can be concluded; claims outside this scope should not be presented as results."
    limits = protocol.get("limitations", "").strip()
    if limits:
        body += f"\n\nDeclared limitation: {limits}"
    return opening + "\n\n" + body


def _mentioned_columns(question: str, df: pd.DataFrame) -> list[str]:
    """Match exact column labels in a command without guessing synonyms."""
    lowered = question.casefold()
    normalised = re.sub(r"[^\w]+", " ", lowered, flags=re.UNICODE)
    matches: list[str] = []
    for column in df.columns:
        label = str(column).casefold()
        simplified = re.sub(r"[^\w]+", " ", label, flags=re.UNICODE).strip()
        if label in lowered or (len(simplified) >= 3 and simplified in normalised):
            matches.append(column)
    return matches


def _method_evidence_table(evidence: pd.DataFrame) -> pd.DataFrame:
    if evidence.empty:
        return pd.DataFrame()
    methods = {
        "OLS / linear regression": r"\bOLS\b|ordinary least squares|linear regression|παλινδρόμη",
        "Panel fixed effects": r"fixed effects?|within estimator|σταθερ(?:ές|ων) επιδρά",
        "Panel random effects": r"random effects?|τυχαί(?:ες|ων) επιδρά",
        "Monte Carlo / simulation": r"monte carlo|simulation|bootstrap|προσομοίω",
        "Clustering": r"cluster(?:ing|s)?|k[- ]?means|hierarchical clustering|ομαδοποίη",
        "Spatial analysis": r"spatial|GIS|Moran(?:'s)? I|χωρικ",
        "MCDA": r"multi.?criteria|MCDA|TOPSIS|PROMETHEE|AHP|πολυκριτήρ",
        "Difference-in-differences": r"difference.?in.?differences?|diff.?in.?diff|DiD",
        "Instrumental variables": r"instrumental variables?|two.?stage least squares|2SLS",
    }
    rows: list[dict[str, Any]] = []
    for row in evidence.itertuples(index=False):
        text = str(row.text)
        for method, pattern in methods.items():
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                start = max(match.start() - 100, 0)
                end = min(match.end() + 180, len(text))
                snippet = re.sub(r"\s+", " ", text[start:end]).strip()
                rows.append({
                    "method": method,
                    "document": row.document,
                    "page": row.page,
                    "evidence_snippet": snippet,
                })
    return pd.DataFrame(rows).drop_duplicates() if rows else pd.DataFrame()


def _result_copy(result: ProtocolResult, algorithm: str) -> ProtocolResult:
    return ProtocolResult(
        algorithm=algorithm,
        tables={name: table.copy() for name, table in result.tables.items()},
        comments=list(result.comments),
        equation=result.equation,
        executed_expression=result.executed_expression,
    )


def _identifier_or_time_like(column: str, series: pd.Series) -> bool:
    """Exclude codes, row identifiers and time co-ordinates from headline findings."""
    name = str(column).casefold()
    identifier_pattern = r"(^|_)(id|code|key|index|serial|mis)(_|$)|a_a_project"
    time_pattern = r"(^|_)(year|date|month|quarter|duration)(_|$)"
    if re.search(identifier_pattern, name) or re.search(time_pattern, name):
        return True
    observed = pd.to_numeric(series, errors="coerce").dropna()
    if len(observed) >= 30 and observed.nunique() / len(observed) >= .98:
        # A nearly unique integer-like column is probably an identifier even if
        # its label is unfamiliar. Continuous measures are not excluded here.
        integer_like = np.allclose(observed, np.round(observed), equal_nan=True)
        if integer_like:
            return True
    return False


def rank_striking_findings(
    df: pd.DataFrame,
    protocol: dict[str, Any],
    current_result: ProtocolResult,
    maximum_variables: int = 30,
) -> tuple[pd.DataFrame, dict[str, Any]]:
    """Rank high-signal findings while suppressing obvious identifiers/artefacts."""
    numeric = [
        c for c in df.select_dtypes(include=np.number).columns
        if not _identifier_or_time_like(c, df[c])
        and df[c].notna().sum() >= max(30, int(.05 * len(df)))
        and df[c].nunique(dropna=True) >= 2
    ][:maximum_variables]
    candidates: list[dict[str, Any]] = []
    for position, left in enumerate(numeric):
        for right in numeric[position + 1:]:
            pair = df[[left, right]].dropna()
            if len(pair) < max(30, int(.05 * len(df))):
                continue
            coefficient, p_value = stats.spearmanr(pair[left], pair[right])
            if not np.isfinite(coefficient):
                continue
            strength = abs(float(coefficient))
            caveat = (
                "Near-overlap: inspect construct definitions and multicollinearity before joint modelling."
                if strength >= .90 else
                "Strong association; direction and causality remain unidentified."
                if strength >= .70 else
                "Moderate association; assess practical significance and robustness."
            )
            candidates.append({
                "finding_type": "Spearman association",
                "variable_1": left,
                "variable_2": right,
                "estimate": float(coefficient),
                "absolute_estimate": strength,
                "n": int(len(pair)),
                "p_value": float(p_value),
                "priority_score": strength,
                "interpretation": caveat,
            })

    ranked = pd.DataFrame(candidates)
    if not ranked.empty:
        p_values = ranked["p_value"].to_numpy(float)
        order = np.argsort(p_values)
        adjusted_sorted = np.minimum.accumulate(
            (p_values[order] * len(p_values) / np.arange(1, len(p_values) + 1))[::-1]
        )[::-1]
        adjusted = np.empty_like(adjusted_sorted)
        adjusted[order] = np.clip(adjusted_sorted, 0, 1)
        ranked["q_value_bh"] = adjusted
        ranked = ranked.sort_values(
            ["priority_score", "n"], ascending=[False, False]
        ).reset_index(drop=True)
        ranked.insert(0, "rank", np.arange(1, len(ranked) + 1))

    missing = _missingness_table(df)
    worst_missing: dict[str, Any] = {}
    if not missing.empty:
        worst = missing.iloc[0]
        worst_missing = {
            "variable": str(worst.variable),
            "missing": int(worst.missing),
            "observed": int(worst.observed),
            "missing_percent": float(worst.missing_percent),
        }

    questionable: list[str] = []
    outcome = protocol.get("outcome")
    if outcome in df and _identifier_or_time_like(outcome, df[outcome]):
        questionable.append(f"the selected outcome ‘{outcome}’ is an identifier/time-type field")
    for predictor in protocol.get("predictors", []) or []:
        if predictor in df and _identifier_or_time_like(predictor, df[predictor]):
            questionable.append(f"the selected predictor ‘{predictor}’ is an identifier/time-type field")

    context = {
        "numeric_candidates": len(numeric),
        "worst_missing": worst_missing,
        "questionable_model_fields": questionable,
        "ols_available": not current_result.tables.get("OLS coefficients", pd.DataFrame()).empty,
    }
    return ranked, context


def _p_text(value: float) -> str:
    if not np.isfinite(value):
        return "not estimable"
    if value < .001:
        return "p < .001"
    return f"p = {value:.3f}"


def _q_text(value: float) -> str:
    if not np.isfinite(value):
        return "BH q not estimable"
    if value < .001:
        return "BH q < .001"
    return f"BH q = {value:.3f}"


def _striking_reply(
    df: pd.DataFrame,
    protocol: dict[str, Any],
    current_result: ProtocolResult,
    ranked: pd.DataFrame,
    context: dict[str, Any],
) -> str:
    records = len(df)
    if ranked.empty:
        return (
            f"Direct answer: no defensible ‘most striking’ association could be ranked in the saved "
            f"scope of {records:,} records. Too few non-identifier numeric measures had sufficient "
            "coverage and variation. Review the Ranked statistical findings and Variable missingness tables before modelling."
        )

    top = ranked.iloc[0]
    direction = "positive" if float(top.estimate) >= 0 else "negative"
    left, right = str(top.variable_1), str(top.variable_2)
    reply = [
        "**Direct answer**",
        (
            f"The strongest statistically defensible pattern in the selected spreadsheet is the "
            f"{direction} association between `{left}` and `{right}`: "
            f"Spearman ρ = {float(top.estimate):.3f}, n = {int(top.n):,}, {_p_text(float(top.p_value))}."
            f" The multiple-screening-adjusted result is {_q_text(float(top.q_value_bh))}."
        ),
        "**Why this is striking**",
        (
            f"An absolute correlation of {abs(float(top.estimate)):.3f} is exceptionally large. "
            "It suggests that the two measures may capture almost the same underlying construct or may be mechanically related. "
            "That makes the pattern important for measurement design and multicollinearity, but not automatically novel or causal."
        ),
    ]
    missing = context.get("worst_missing") or {}
    if missing and float(missing.get("missing_percent", 0)) > 20:
        reply.extend([
            "**Most important data-quality warning**",
            (
                f"`{missing['variable']}` is {float(missing['missing_percent']):.2f}% missing "
                f"({int(missing['observed']):,} observed and {int(missing['missing']):,} missing values). "
                "It should not be used as a principal outcome without recovery or a defensible missing-data strategy."
            ),
        ])
    questionable = context.get("questionable_model_fields") or []
    if questionable:
        reply.extend([
            "**Model-specification warning**",
            "The currently selected OLS should not be reported as a substantive model because " + " and ".join(questionable) + ". Statistical significance cannot rescue a scientifically meaningless specification.",
        ])
    reply.extend([
        "**Publication-ready wording**",
        (
            f"“Across {int(top.n):,} projects, `{left}` was very strongly and {direction}ly associated with `{right}` "
            f"(Spearman ρ = {float(top.estimate):.3f}, {_p_text(float(top.p_value))}, {_q_text(float(top.q_value_bh))}). "
            "The magnitude indicates potential construct overlap; accordingly, the indicators were not interpreted as independent causal determinants and require collinearity and operational-definition checks before inclusion in the same model.”"
        ),
        "**What to do next**",
        "Inspect the top ten rows of the new Ranked statistical findings table, verify the definitions of the leading pair, then run VIF/PCA or retain only the theoretically primary indicator. Use the dedicated econometric module only after selecting a substantively meaningful outcome and predictors.",
    ])
    return "\n\n".join(reply)


def _limitations_table(df: pd.DataFrame, protocol: dict[str, Any]) -> pd.DataFrame:
    missing = _missingness_table(df)
    worst = missing.iloc[0] if not missing.empty else None
    rows = [
        {
            "limitation_area": "Study design and identification",
            "evidence_in_current_scope": "No random assignment or explicit quasi-experimental identification is documented in the saved protocol.",
            "why_it_matters": "Observed relationships may reflect confounding, selection or reverse causality.",
            "paper_ready_statement": "The observational design supports descriptive and conditional associative interpretation but does not, by itself, identify causal effects.",
        },
        {
            "limitation_area": "Sample and external validity",
            "evidence_in_current_scope": f"The analytical scope contains {len(df):,} retained records after the selected filters.",
            "why_it_matters": "The retained projects may not represent other programmes, periods or populations.",
            "paper_ready_statement": "Generalisability is bounded by the programme coverage, inclusion rules and time period represented in the analytical sample.",
        },
        {
            "limitation_area": "Measurement",
            "evidence_in_current_scope": "Variable labels and administrative indicators may overlap or reflect programme reporting rules rather than independent constructs.",
            "why_it_matters": "Overlapping indicators can inflate correlations and destabilise regression coefficients.",
            "paper_ready_statement": "Several administrative indicators may capture related constructs; operational definitions and collinearity were therefore examined before joint modelling.",
        },
        {
            "limitation_area": "Multiple testing",
            "evidence_in_current_scope": "Exploratory screening can compare many variables and models.",
            "why_it_matters": "Some small p-values can arise by chance when many hypotheses are tested.",
            "paper_ready_statement": "Exploratory pairwise results were treated as hypothesis-generating and adjusted for the false-discovery burden where applicable.",
        },
    ]
    if worst is not None and float(worst.missing_percent) > 0:
        rows.insert(1, {
            "limitation_area": "Missing data",
            "evidence_in_current_scope": f"{worst.variable} is {float(worst.missing_percent):.2f}% missing ({int(worst.observed):,} observed values).",
            "why_it_matters": "Sparse variables can reduce precision, alter the effective sample and bias estimates if missingness is systematic.",
            "paper_ready_statement": "Results involving incomplete variables are conditional on observed cases and may be sensitive to the mechanism generating missingness.",
        })
    return pd.DataFrame(rows)


def _claim_matrix(df: pd.DataFrame, protocol: dict[str, Any], result: ProtocolResult) -> pd.DataFrame:
    ranked, context = rank_striking_findings(df, protocol, result)
    rows: list[dict[str, str]] = []
    if not ranked.empty:
        top = ranked.iloc[0]
        rows.append({
            "claim_status": "Safe with qualification",
            "claim": f"{top.variable_1} and {top.variable_2} are strongly associated in the selected sample (Spearman rho={float(top.estimate):.3f}, n={int(top.n):,}).",
            "reason": "This statement reports the observed association, sample and magnitude without assigning causality.",
        })
    rows.extend([
        {"claim_status": "Safe", "claim": f"The reported outputs describe the selected analytical scope of {len(df):,} records.", "reason": "This is a direct sample statement."},
        {"claim_status": "Not safe without identification", "claim": "A predictor caused the observed outcome.", "reason": "Regression/correlation alone does not eliminate confounding, selection or reverse causality."},
        {"claim_status": "Not safe", "claim": "A statistically significant result is necessarily important or policy-effective.", "reason": "Statistical significance does not measure practical magnitude, validity or implementation value."},
        {"claim_status": "Not safe", "claim": "The findings automatically generalise to other programmes, countries or periods.", "reason": "External validity is bounded by the observed sample and data-generating context."},
    ])
    if context.get("questionable_model_fields"):
        rows.append({
            "claim_status": "Not safe",
            "claim": "The currently selected OLS is a substantive explanatory model.",
            "reason": "; ".join(context["questionable_model_fields"]) + ".",
        })
    return pd.DataFrame(rows)


def _further_analysis_table(df: pd.DataFrame, protocol: dict[str, Any], result: ProtocolResult) -> pd.DataFrame:
    year = protocol.get("year_column") if protocol.get("year_column") in df else next((c for c in df if "year" in str(c).casefold()), None)
    region = next((c for c in df if "region" in str(c).casefold() or "nuts" in str(c).casefold()), None)
    # The prioritisation is intentionally explicit rather than pretending every
    # possible method is appropriate for every dataset.
    rows = [
        {"priority": 1, "analysis": "Measurement overlap and multicollinearity audit", "scientific_value": "High", "feasibility": "Immediate", "requires": "Current numeric indicators", "why": "Very strong correlations may indicate redundant constructs; use definitions, VIF and PCA before multivariable modelling."},
        {"priority": 2, "analysis": "Theory-led robust regression", "scientific_value": "High", "feasibility": "Immediate after variable selection", "requires": "A meaningful outcome, predictors and controls", "why": "Moves from unadjusted patterns to conditional estimates with uncertainty and diagnostics."},
        {"priority": 3, "analysis": "Monte Carlo / bootstrap stability", "scientific_value": "Medium–high", "feasibility": "Immediate after OLS", "requires": "A valid OLS specification, seed and repetitions", "why": "Shows whether coefficient signs and intervals are stable under resampling assumptions."},
        {"priority": 4, "analysis": "Cross-validated prediction", "scientific_value": "Medium–high", "feasibility": "Immediate", "requires": "Continuous outcome and non-leaking predictors", "why": "Tests generalisation and identifies predictors that improve out-of-sample accuracy."},
        {"priority": 5, "analysis": "Cluster/typology validation", "scientific_value": "Medium", "feasibility": "Immediate", "requires": "Standardised substantive indicators", "why": "Identifies heterogeneous project profiles and tests whether absorption-only clusters persist in richer specifications."},
    ]
    if year:
        rows.insert(2, {"priority": 3, "analysis": "Longitudinal/panel robustness", "scientific_value": "High", "feasibility": "Depends on repeated entities", "requires": f"Entity identifier plus {year}", "why": "Separates within-entity temporal change from stable between-entity differences."})
    if region:
        rows.append({"priority": len(rows) + 1, "analysis": "Spatial clustering and spatial regression", "scientific_value": "Medium–high", "feasibility": "Immediate for regional aggregates", "requires": f"{region}, defensible rate/measure and spatial weights", "why": "Tests whether regional patterns cluster geographically rather than treating mapped differences as independent."})
    for index, row in enumerate(rows, start=1):
        row["priority"] = index
    return pd.DataFrame(rows)


def _research_question_table(df: pd.DataFrame, protocol: dict[str, Any], result: ProtocolResult) -> pd.DataFrame:
    ranked, _ = rank_striking_findings(df, protocol, result)
    rows: list[dict[str, Any]] = []
    for item in ranked.head(6).itertuples(index=False):
        direction = "positive" if float(item.estimate) >= 0 else "inverse"
        rows.append({
            "research_question": f"How strongly are {item.variable_1} and {item.variable_2} associated in the selected sample?",
            "recommended_method": "Spearman correlation with Benjamini–Hochberg correction",
            "current_answer_status": "Answered from current results",
            "numerical_evidence": f"rho={float(item.estimate):.3f}; n={int(item.n):,}; {_p_text(float(item.p_value))}; {_q_text(float(item.q_value_bh))}",
            "safe_answer": f"The indicators have a {direction} monotonic association of very {'large' if abs(float(item.estimate)) >= .7 else 'moderate' if abs(float(item.estimate)) >= .4 else 'limited'} magnitude. This is an observed relationship, not a causal effect.",
            "remaining_uncertainty": "Construct overlap, confounding, measurement rules and temporal ordering have not been resolved by the bivariate test.",
        })
    missing = _missingness_table(df)
    if not missing.empty and float(missing.iloc[0].missing_percent) > 0:
        item = missing.iloc[0]
        rows.append({
            "research_question": f"Is {item.variable} sufficiently complete for primary modelling?",
            "recommended_method": "Variable-level missingness audit and missing-data mechanism assessment",
            "current_answer_status": "Answered for coverage; mechanism remains unresolved",
            "numerical_evidence": f"Observed={int(item.observed):,}; missing={int(item.missing):,}; missing={float(item.missing_percent):.2f}%",
            "safe_answer": "Its present coverage is insufficient for an ordinary primary model unless the data can be recovered or a defensible sparse-data strategy is justified.",
            "remaining_uncertainty": "The reason values are absent and whether missingness depends on observed or unobserved quantities require source investigation.",
        })
    if not rows:
        rows.append({
            "research_question": "Which defensible statistical questions can the retained evidence answer?",
            "recommended_method": "Data-quality and measurement review before modelling",
            "current_answer_status": "Not yet answerable",
            "numerical_evidence": "No adequately covered non-identifier numeric pair was available.",
            "safe_answer": "The current scope supports data description only.",
            "remaining_uncertainty": "Variable definitions, coverage and analytical roles must be resolved.",
        })
    return pd.DataFrame(rows)


def _methodology_outline_table(
    df: pd.DataFrame,
    protocol: dict[str, Any],
    result: ProtocolResult,
) -> pd.DataFrame:
    """Return a downloadable, paper-ordered methods checklist."""
    profile = scope_profile(df).iloc[0]
    outcome = protocol.get("outcome") or "Not yet selected"
    predictors = protocol.get("predictors") or []
    predictor_text = ", ".join(map(str, predictors)) if predictors else "Not yet selected"
    return pd.DataFrame([
        {
            "step": 1,
            "methods_subsection": "Design and unit of analysis",
            "current_setting": f"{int(profile.records):,} retained rows; confirm what one row represents.",
            "what_to_write": "Name the study design, population, programme context and unit represented by each row.",
            "why_it_matters": "Readers cannot interpret estimates unless the observational unit and design are explicit.",
        },
        {
            "step": 2,
            "methods_subsection": "Sample construction",
            "current_setting": f"Year range: {protocol.get('year_range')}; filters: {protocol.get('filters') or 'none recorded'}.",
            "what_to_write": "Report inclusion, exclusion, year, file, sheet, join and duplicate-handling rules.",
            "why_it_matters": "These choices define which observations the findings describe.",
        },
        {
            "step": 3,
            "methods_subsection": "Variables and operationalisation",
            "current_setting": f"Outcome: {outcome}; predictors: {predictor_text}.",
            "what_to_write": "Define every variable, unit, denominator, transformation, reference group and expected sign.",
            "why_it_matters": "A coefficient has no scientific meaning without a valid construct and measurement scale.",
        },
        {
            "step": 4,
            "methods_subsection": "Model and estimand",
            "current_setting": f"Algorithm: {result.algorithm}; equation: {protocol.get('equation') or 'not specified'}.",
            "what_to_write": "State the exact equation, estimator, target quantity and interpretation of every symbol.",
            "why_it_matters": "This identifies the question that the numerical estimate actually answers.",
        },
        {
            "step": 5,
            "methods_subsection": "Uncertainty and diagnostics",
            "current_setting": "Use the covariance, confidence-level, residual, fit, VIF and influence outputs produced by the selected module.",
            "what_to_write": "Report standard-error or resampling choices, assumptions, diagnostics and complete-case sample size.",
            "why_it_matters": "Uncertainty and model checks determine how much weight the estimate deserves.",
        },
        {
            "step": 6,
            "methods_subsection": "Robustness and validation",
            "current_setting": "Pre-specify defensible alternative samples, models, temporal windows and simulation settings.",
            "what_to_write": "Explain which alternatives were tested, why they are defensible and whether conclusions changed.",
            "why_it_matters": "Robustness distinguishes a persistent finding from a result tied to one convenient specification.",
        },
        {
            "step": 7,
            "methods_subsection": "Reproducibility",
            "current_setting": "The Research Chair bundle stores the filtered data, protocol, evidence, tables, seed and prompt library.",
            "what_to_write": "Archive the analytical sample, settings, software version, seed, repetitions and generated outputs.",
            "why_it_matters": "Another researcher should be able to reconstruct the reported analysis.",
        },
    ])


def _command_reply(
    action: str,
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
) -> str:
    profile = result.tables.get("Scope profile", pd.DataFrame())
    row = profile.iloc[0] if not profile.empty else None
    records = int(row.records) if row is not None else 0
    variables = int(row.variables) if row is not None else 0
    numeric_variables = int(row.numeric_variables) if row is not None else 0
    table_descriptions = [
        f"{name} ({len(table):,} rows)"
        for name, table in result.tables.items() if table is not None and not table.empty
    ]
    opening = (
        f"I executed a data-grounded {action} on the saved analytical scope: "
        f"{records:,} records, {variables:,} variables and {numeric_variables:,} numeric variables."
    )
    details: list[str] = []
    correlations = result.tables.get("Correlation screening", pd.DataFrame())
    if not correlations.empty:
        strongest = correlations.iloc[0]
        details.append(
            f"The strongest screened monotonic association is {strongest.variable_1} versus "
            f"{strongest.variable_2} (Spearman ρ={float(strongest.spearman_rho):.3f}, "
            f"n={int(strongest.n):,}, p={float(strongest.p_value):.4g})."
        )
    trend = result.tables.get("Longitudinal results", pd.DataFrame())
    if not trend.empty and "year" in trend:
        years = pd.to_numeric(trend["year"], errors="coerce").dropna()
        if not years.empty:
            details.append(
                f"The longitudinal output covers {int(years.min())}–{int(years.max())} "
                f"across {years.nunique()} observed years."
            )
    fit = result.tables.get("OLS fit", pd.DataFrame())
    if not fit.empty:
        details.append(
            f"The HC3 OLS specification uses n={int(fit.iloc[0].n):,} complete observations "
            f"and has adjusted R²={float(fit.iloc[0].adjusted_r_squared):.3f}."
        )
    methods = result.tables.get("Methods found in selected PDFs", pd.DataFrame())
    if not methods.empty:
        details.append(
            "The selected PDF pages explicitly mention: "
            + ", ".join(methods.method.drop_duplicates().head(8)) + "."
        )
    produced = "Generated result tables: " + "; ".join(table_descriptions) + "."
    safeguards = (
        "These are computed results from the selected XLSX/CSV rows and columns. "
        "Correlations and temporal movements are not causal effects; a dependent variable and predictors "
        "must be explicitly selected before the app will estimate OLS."
    )
    limitations = str(protocol.get("limitations", "")).strip()
    if limitations:
        safeguards += f" Declared limitation: {limitations}"
    return "\n\n".join([opening, produced] + details + [safeguards])


def classify_command_feasibility(
    question: str,
    df: pd.DataFrame,
    protocol: dict[str, Any],
    evidence: pd.DataFrame,
) -> FeasibilityDecision:
    """Classify a command before execution without pretending unsupported work ran."""
    q = question.casefold().strip()
    if not q:
        return FeasibilityDecision(
            "NOT FEASIBLE — LOGIC/SETUP ERROR",
            "The question is blank. A blank instruction has no estimand, operation or requested output.",
            "Edit the question before running it.",
            False,
        )
    if df.empty and not any(term in q for term in ["pdf", "paper", "methodology", "equation", "algorithm"]):
        return FeasibilityDecision(
            "NOT FEASIBLE — LOGIC/SETUP ERROR",
            "The command requests statistical analysis but the saved analytical scope contains no spreadsheet rows.",
            "Upload or activate an XLSX/CSV dataset, then run the Chair again.",
            False,
        )

    invalid_patterns = {
        "prove causality": "An observational calculation cannot prove causality without a defensible identification design.",
        "prove that .* causes": "The requested causal conclusion is not identified by correlation or ordinary regression.",
        "divide by zero": "Division by zero is undefined.",
        "p[- ]?value (?:above|greater than) 1": "A probability and therefore a p-value must lie between 0 and 1.",
        "execute .*python": "Arbitrary code execution is deliberately blocked; only validated analytical routines may run.",
        "ignore .*assumption": "A method cannot be made valid by instructing the system to ignore its assumptions.",
    }
    for pattern, reason in invalid_patterns.items():
        if re.search(pattern, q):
            return FeasibilityDecision(
                "NOT FEASIBLE — STATISTICAL/MATHEMATICAL ERROR",
                reason,
                "Reformulate the question around a valid estimand, design and implemented estimator.",
                False,
            )

    outcome = protocol.get("outcome") if protocol.get("outcome") in df else None
    predictors = [c for c in (protocol.get("predictors") or []) if c in df and c != outcome]
    year = protocol.get("year_column") if protocol.get("year_column") in df else None
    group = protocol.get("group_column") if protocol.get("group_column") in df else None

    requirements = [
        (("monte carlo", "bootstrap simulation"), bool(outcome and predictors), "Monte Carlo OLS requires one explicit outcome and at least one predictor."),
        (("regression", "ols", "coefficient"), bool(outcome and predictors), "Regression requires one explicit outcome and at least one predictor."),
        (("trend", "through the years", "over time", "time series", "arima", "forecast"), bool(year), "A longitudinal or forecasting command requires a valid year/time variable."),
        (("group test", "differs across", "difference across", "anova", "mann-whitney", "kruskal"), bool(outcome and group), "A group comparison requires a numeric outcome and a grouping variable."),
        (("pdf methodology", "methodology of this paper", "copy the methodology", "extract the methodology"), not evidence.empty, "PDF methodology extraction requires at least one selected PDF page."),
        (("predict", "cross-validation", "random forest", "elastic net"), bool(outcome and predictors), "Prediction requires a continuous outcome and at least one non-leaking predictor."),
        (("panel", "fixed effects", "random effects", "hausman"), bool(outcome and predictors and year and group), "Panel estimation requires entity, time, outcome and predictor mappings."),
    ]
    for terms, ready, reason in requirements:
        if any(term in q for term in terms) and not ready:
            return FeasibilityDecision(
                "NOT FEASIBLE — LOGIC/SETUP ERROR",
                reason,
                "Map the required variables in Quick model settings; the Chair will not invent them.",
                False,
            )

    directly_supported = [
        "striking", "important result", "important facts", "key facts", "executive summary", "strongest result", "main finding", "headline",
        "missing", "quality", "duplicate", "correlation", "association", "relationship",
        "trend", "through the years", "over time", "year", "regression", "ols", "coefficient",
        "monte carlo", "bootstrap", "conclude", "conclusion", "limitation", "further analysis", "additional analysis", "additional analyses",
        "further research", "research question", "methodology outline", "outline of the methodology",
        "results outline", "outline of the results", "paper", "article", "publish", "pdf",
        "equation", "algorithm", "normality", "outlier", "group test", "differs across",
        "pca", "principal component", "cronbach", "reliability", "cluster", "segmentation",
        "predict", "cross-validation", "time series", "arima", "forecast", "panel", "fixed effects",
    ]
    if any(term in q for term in directly_supported):
        return FeasibilityDecision(
            "FEASIBLE — EXECUTED",
            "The requested method is implemented and the required evidence/settings are present.",
            "The Research Chair computed the result from the saved analytical scope.",
            True,
        )

    menu_only = {
        "instrumental variable": "IV/2SLS instrument and endogenous-variable mapping",
        "2sls": "IV/2SLS instrument and endogenous-variable mapping",
        "difference-in-differences": "treatment/post/control mapping",
        "difference in differences": "treatment/post/control mapping",
        "moran": "spatial weights and geography mapping",
        "lisa": "spatial weights and geography mapping",
        "gis": "geography-level mapping",
        "mcda": "criterion directions and weights",
        "topsis": "criterion directions and weights",
        "promethee": "criterion directions and weights",
        "ahp": "pairwise judgement matrix",
        "allocation": "cost, benefit, budget and constraint mappings",
        "1000 × 1000": "outcome/predictor matrix configuration",
    }
    for term, missing_route in menu_only.items():
        if term in q:
            return FeasibilityDecision(
                "NOT FEASIBLE — NOT YET IMPLEMENTED IN CHAT",
                f"The original menu contains this analytical family, but its {missing_route} is not yet exposed safely through the chat substitute.",
                "Use the original specialist module or contact the developers to add a validated chat runner.",
                False,
            )

    return FeasibilityDecision(
        "NOT FEASIBLE — NOT YET IMPLEMENTED IN CHAT",
        "The deterministic free interpreter cannot map this wording to a validated analytical routine without guessing.",
        "Contact the developers to add the requested method or wording to the Research Chair command map.",
        False,
    )


def execute_question_batch(
    df: pd.DataFrame,
    questions: list[str],
    protocol: dict[str, Any],
    current_result: ProtocolResult,
    evidence: pd.DataFrame,
) -> tuple[ProtocolResult, pd.DataFrame, str]:
    """Run selected editable questions and preserve every verdict and answer."""
    combined = _result_copy(current_result, "Guided autopilot question batch")
    rows: list[dict[str, Any]] = []
    prose: list[str] = ["# Research Chair — genuine answers to selected questions"]
    for number, raw_question in enumerate(questions, start=1):
        question = raw_question.strip()
        decision = classify_command_feasibility(question, df, protocol, evidence)
        answer = f"**{decision.reason}**\n\n{decision.route}"
        if decision.can_execute:
            try:
                executed, answer = execute_natural_language_command(df, question, protocol, combined, evidence)
                for name, table in executed.tables.items():
                    if table is None:
                        continue
                    export_name = name
                    if export_name in combined.tables and not combined.tables[export_name].equals(table):
                        export_name = f"Q{number:02d} — {name}"
                    combined.tables[export_name] = table.copy()
                combined.comments.extend(c for c in executed.comments if c not in combined.comments)
            except Exception as exc:
                decision = FeasibilityDecision(
                    "NOT FEASIBLE — LOGIC/SETUP ERROR",
                    f"The validated routine stopped before producing results: {exc}",
                    "Correct the named data or model requirement and run the question again.",
                    False,
                )
                answer = f"**{decision.reason}**\n\n{decision.route}"
        rows.append({
            "question_number": number,
            "question": question,
            "feasibility_verdict": decision.status,
            "reason": decision.reason,
            "route": decision.route,
            "answer": answer,
        })
        prose.extend([
            "",
            f"## Question {number}: {question}",
            "",
            f"**Feasibility verdict:** {decision.status}",
            "",
            answer,
        ])
    answer_table = pd.DataFrame(rows)
    combined.tables["Question answers and feasibility"] = answer_table
    return combined, answer_table, "\n".join(prose).strip() + "\n"


def execute_natural_language_command(
    df: pd.DataFrame,
    question: str,
    protocol: dict[str, Any],
    current_result: ProtocolResult,
    evidence: pd.DataFrame,
) -> tuple[ProtocolResult, str]:
    """Execute bounded statistical commands against the already-saved scope.

    This is deliberately deterministic and free: it identifies an analytical
    intent, computes tables from the data, and then describes only those tables.
    It never treats free-form prose as arbitrary Python or silently invents a
    dependent variable.
    """
    if df.empty:
        return current_result, "No spreadsheet rows are available in the saved scope. Upload/select XLSX or CSV evidence and run the Research Command first."

    q = question.casefold().strip()
    mentioned = _mentioned_columns(question, df)
    numeric_all = list(df.select_dtypes(include=np.number).columns)
    mentioned_numeric = [c for c in mentioned if c in numeric_all]
    outcome = protocol.get("outcome") if protocol.get("outcome") in df else None
    predictors = [c for c in protocol.get("predictors", []) if c in df and c != outcome]
    year_column = protocol.get("year_column") if protocol.get("year_column") in df else None
    group_column = protocol.get("group_column") if protocol.get("group_column") in df else None
    aggregation = protocol.get("aggregation") or "Mean"
    equation = protocol.get("equation", "")

    paper_intent = any(term in q for term in [
        "run the analysis", "as in paper", "paper", "article", "publish", "full analysis",
        "γράψ", "δημοσί", "ανάλυση όπως", "τρέξε την ανάλυση", "εργασία",
    ])
    regression_intent = any(term in q for term in ["regression", "ols", "coefficient", "παλινδ", "συντελεστ"])
    trend_intent = any(term in q for term in ["trend", "through the years", "over time", "year", "χρον", "έτος", "ετ"])
    correlation_intent = any(term in q for term in ["correlation", "association", "relationship", "συσχέτ", "σχέση"])
    quality_intent = any(term in q for term in ["missing", "quality", "clean", "duplicate", "ελλιπ", "ποιότη", "διπλότυπ"])
    pdf_intent = any(term in q for term in ["pdf", "note", "literature", "method in", "θεωρ", "βιβλιο", "σημειώ"])
    striking_intent = any(term in q for term in [
        "most striking", "most important result", "strongest result", "main finding",
        "headline result", "most remarkable", "key finding", "πιο εντυπωσιακ",
        "σημαντικότερο αποτέλεσμα", "ισχυρότερο αποτέλεσμα", "κύριο εύρημα",
    ])
    executive_intent = any(term in q for term in ["five most important facts", "most important facts", "executive summary", "key facts", "principal facts"])
    monte_carlo_intent = any(term in q for term in ["monte carlo", "bootstrap simulation", "προσομοίωση monte", "μοντε καρλο"])
    conclusion_intent = any(term in q for term in ["safely conclude", "cannot conclude", "can i conclude", "what can i conclude", "can i not conclude", "what can i not conclude", "safe conclusion", "τι μπορώ να συμπεράνω", "δεν μπορώ να συμπεράνω"])
    limitation_intent = any(term in q for term in ["limitation", "limitations section", "περιορισμ", "αδυναμί"])
    further_intent = any(term in q for term in ["further analysis", "additional analysis", "additional analyses", "further research", "future research", "next analysis", "επόμενη ανάλυση", "περαιτέρω έρευνα"])
    research_question_intent = any(term in q for term in ["research questions and", "propose research questions", "generate research questions", "ερευνητικά ερωτήματα", "research question proposals"])
    hypothesis_intent = any(term in q for term in ["testable hypotheses", "convert my principal research question", "generate hypotheses", "hypothesis and operationalisation"])
    methodology_outline_intent = any(term in q for term in ["outline of the methodology", "methodology outline", "methodology used", "methods outline", "περίγραμμα μεθοδολογ"])
    results_outline_intent = any(term in q for term in ["outline of the results", "results outline", "result section outline", "περίγραμμα αποτελεσμ"])
    pdf_methodology_intent = any(term in q for term in ["copy the methodology", "replicate the methodology", "methodology of this paper", "extract the methodology", "μεθοδολογία του άρθρου", "αντιγραφή μεθοδολογ"])
    equation_audit_intent = any(term in q for term in ["check this equation", "audit this equation", "equation correctness", "execute this equation", "check this algorithm", "audit this algorithm", "execute this algorithm", "mathematical correctness"])
    normality_intent = any(term in q for term in ["normality", "normally distributed", "shapiro", "d'agostino", "anderson-darling"])
    outlier_intent = any(term in q for term in ["outlier", "extreme value", "unusual value"])
    group_test_intent = any(term in q for term in ["group test", "differs across", "difference across", "anova", "mann-whitney", "kruskal"])
    pca_intent = any(term in q for term in ["pca", "principal component"])
    reliability_intent = any(term in q for term in ["cronbach", "reliability", "combined scale"])
    clustering_intent = any(term in q for term in ["cluster", "segmentation", "typolog"])
    prediction_intent = any(term in q for term in ["cross-validation", "cross validated", "predictive model", "random forest", "elastic net"])
    forecast_intent = any(term in q for term in ["arima", "forecast"])
    time_series_intent = any(term in q for term in ["time series", "stationarity", "unit root", "adf", "kpss"])
    panel_intent = any(term in q for term in ["panel", "fixed effects", "random effects", "hausman"])

    if executive_intent:
        ranked, context = rank_striking_findings(df, protocol, current_result)
        profile = scope_profile(df).iloc[0]
        missing = _missingness_table(df)
        facts: list[dict[str, Any]] = [
            {"rank": 1, "finding": "Analytical scope", "numerical_evidence": f"{int(profile.records):,} records; {int(profile.variables):,} variables; {int(profile.numeric_variables):,} numeric", "plain_language_meaning": "This is the exact population of rows and columns described by the current answers.", "paper_ready_sentence": f"The analytical sample comprised {int(profile.records):,} records across {int(profile.variables):,} retained variables, of which {int(profile.numeric_variables):,} were numeric."},
            {"rank": 2, "finding": "Overall completeness", "numerical_evidence": f"{int(profile.missing_cells):,} missing cells ({float(profile.missing_percent):.2f}%); {int(profile.duplicate_rows):,} duplicate rows", "plain_language_meaning": "The aggregate percentage can conceal one almost-empty variable, so variable-level coverage must also be examined.", "paper_ready_sentence": f"Across the retained matrix, {float(profile.missing_percent):.2f}% of cells were missing and {int(profile.duplicate_rows):,} exact duplicate rows were detected."},
        ]
        if not ranked.empty:
            top = ranked.iloc[0]
            facts.append({"rank": 3, "finding": "Strongest defensible association", "numerical_evidence": f"{top.variable_1} ↔ {top.variable_2}: rho={float(top.estimate):.3f}; n={int(top.n):,}; {_q_text(float(top.q_value_bh))}", "plain_language_meaning": "The two measures move together very strongly; this may reveal common content or a mechanical link.", "paper_ready_sentence": f"The strongest screened association linked {top.variable_1} and {top.variable_2} (Spearman rho={float(top.estimate):.3f}, n={int(top.n):,}, {_q_text(float(top.q_value_bh))})."})
        if not missing.empty:
            worst = missing.iloc[0]
            facts.append({"rank": 4, "finding": "Largest data-quality constraint", "numerical_evidence": f"{worst.variable}: {float(worst.missing_percent):.2f}% missing; {int(worst.observed):,} observed", "plain_language_meaning": "This variable cannot support an ordinary primary model at its present coverage.", "paper_ready_sentence": f"The most severe coverage limitation concerned {worst.variable}, for which {float(worst.missing_percent):.2f}% of observations were absent."})
        numeric = [c for c in df.select_dtypes(include=np.number) if not _identifier_or_time_like(c, df[c])]
        if numeric:
            variation = pd.DataFrame([{"variable": c, "coefficient_of_variation": abs(float(pd.to_numeric(df[c], errors='coerce').std() / pd.to_numeric(df[c], errors='coerce').mean())) if pd.to_numeric(df[c], errors='coerce').mean() not in [0, np.nan] else np.nan} for c in numeric]).replace([np.inf, -np.inf], np.nan).dropna()
            if not variation.empty:
                item = variation.sort_values("coefficient_of_variation", ascending=False).iloc[0]
                facts.append({"rank": 5, "finding": "Greatest relative dispersion", "numerical_evidence": f"{item.variable}: coefficient of variation={float(item.coefficient_of_variation):.3f}", "plain_language_meaning": "Observed values are highly heterogeneous relative to their mean, making the median and distribution more informative than the mean alone.", "paper_ready_sentence": f"Relative dispersion was greatest for {item.variable} (coefficient of variation={float(item.coefficient_of_variation):.3f})."})
        facts_table = pd.DataFrame(facts).head(5)
        result = _result_copy(current_result, "Executive analytical synthesis")
        result.tables["Five principal findings"] = facts_table
        reply = "**Five evidence-grounded facts**\n\n" + "\n\n".join(
            f"{int(row.rank)}. **{row.finding}.** {row.numerical_evidence}. {row.plain_language_meaning}\n\n**Draft paper sentence:** {row.paper_ready_sentence}"
            for row in facts_table.itertuples(index=False)
        )
        reply += "\n\nThe association and dispersion findings are exploratory; neither should be described as a causal effect."
        return result, reply

    if striking_intent:
        ranked, context = rank_striking_findings(df, protocol, current_result)
        result = _result_copy(current_result, "Ranked headline-finding analysis")
        result.tables["Scope profile"] = scope_profile(df)
        result.tables["Descriptive statistics"] = _descriptive_table(df, list(df.select_dtypes(include=np.number).columns))
        result.tables["Variable missingness"] = _missingness_table(df)
        result.tables["Ranked statistical findings"] = ranked.head(100)
        result.comments.append(
            "Headline findings exclude apparent identifiers, codes and time co-ordinates and require adequate coverage and variation."
        )
        result.comments.append(
            "Ranking is an exploratory salience screen, not a substitute for theory, identification, multiplicity control or external validation."
        )
        return result, _striking_reply(df, protocol, current_result, ranked, context)

    if conclusion_intent:
        claims = _claim_matrix(df, protocol, current_result)
        result = _result_copy(current_result, "Permitted-claim assessment")
        result.tables["Permitted and prohibited claims"] = claims
        safe = claims[claims.claim_status.str.startswith("Safe")]
        unsafe = claims[claims.claim_status.str.startswith("Not safe")]
        reply = "**What can be concluded safely**\n\n" + "\n".join(f"- {row.claim} Reason: {row.reason}" for row in safe.itertuples())
        reply += "\n\n**What cannot be concluded**\n\n" + "\n".join(f"- {row.claim} Reason: {row.reason}" for row in unsafe.itertuples())
        reply += "\n\nThe complete permitted/prohibited claim matrix has been added to the downloadable results."
        return result, reply

    if limitation_intent:
        limitations_table = _limitations_table(df, protocol)
        result = _result_copy(current_result, "Evidence-grounded limitations assessment")
        result.tables["Paper limitations"] = limitations_table
        paragraphs = "\n\n".join(str(value) for value in limitations_table.paper_ready_statement)
        reply = "**Paper-ready limitations section**\n\n" + paragraphs
        reply += "\n\nEach statement is tied to the current sample or design and is included in the new Paper limitations table."
        return result, reply

    if further_intent:
        roadmap = _further_analysis_table(df, protocol, current_result)
        result = _result_copy(current_result, "Prioritised further-analysis roadmap")
        result.tables["Further analysis roadmap"] = roadmap
        recommendations = "\n".join(
            f"{int(row.priority)}. **{row.analysis}** — {row.why} Required: {row.requires}"
            for row in roadmap.itertuples()
        )
        reply = "**Prioritised further analysis and research agenda**\n\n" + recommendations
        reply += "\n\nThe ranking separates analyses feasible with the current data from those that depend on a valid outcome, repeated entities or spatial structure."
        return result, reply

    if hypothesis_intent:
        ranked, _ = rank_striking_findings(df, protocol, current_result)
        rows: list[dict[str, Any]] = []
        for index, item in enumerate(ranked.head(6).itertuples(index=False), start=1):
            expected = "positive" if float(item.estimate) >= 0 else "negative"
            rows.append({
                "hypothesis": f"H{index}",
                "testable_statement": f"{item.variable_1} is {expected}ly associated with {item.variable_2} in the selected analytical population.",
                "variables": f"{item.variable_1}; {item.variable_2}",
                "operational_test": "Spearman rank correlation with BH false-discovery adjustment",
                "current_evidence": f"rho={float(item.estimate):.3f}; n={int(item.n):,}; {_q_text(float(item.q_value_bh))}",
                "interpretive_boundary": "Associative hypothesis only; causal wording requires an identification design.",
            })
        table = pd.DataFrame(rows)
        result = _result_copy(current_result, "Testable hypothesis development")
        result.tables["Hypotheses and operationalisation"] = table
        if table.empty:
            return result, "**Not feasible with the current evidence.** No adequately covered non-identifier numeric pairs were available from which to formulate data-grounded hypotheses."
        reply = "**Data-grounded, testable hypotheses**\n\n" + "\n\n".join(
            f"**{row.hypothesis}.** {row.testable_statement} **Operational test:** {row.operational_test}. **Current evidence:** {row.current_evidence}. **Boundary:** {row.interpretive_boundary}"
            for row in table.itertuples(index=False)
        )
        reply += "\n\nBecause these hypotheses were generated after inspecting the same data, they are exploratory and should be confirmed in a new sample or a pre-specified holdout."
        return result, reply

    if research_question_intent:
        questions = _research_question_table(df, protocol, current_result)
        result = _result_copy(current_result, "Research-question and answer development")
        result.tables["Research question proposals"] = questions
        reply = "**Research questions with genuine answers from the current results**\n\n" + "\n\n".join(
            f"{index}. **Question:** {row.research_question}\n\n**Answer:** {row.safe_answer}\n\n**Numerical evidence:** {row.numerical_evidence}. **Method:** {row.recommended_method}. **Remaining uncertainty:** {row.remaining_uncertainty}"
            for index, row in enumerate(questions.itertuples(), start=1)
        )
        return result, reply

    if methodology_outline_intent:
        scope = scope_profile(df).iloc[0]
        methodology = _methodology_outline_table(df, protocol, current_result)
        result = _result_copy(current_result, "Methods-section outline")
        result.tables["Methodology outline"] = methodology
        missing_pct = float(scope.missing_percent)
        reply = (
            "**Draft methodology text**\n\n"
            f"The study used an observational quantitative design based on the saved analytical dataset. The final analytical scope comprised {int(scope.records):,} records and {int(scope.variables):,} retained variables, including {int(scope.numeric_variables):,} numeric measures. Exact inclusion was determined by the saved file, column and row selections; {missing_pct:.2f}% of cells were incomplete and {int(scope.duplicate_rows):,} exact duplicate rows were identified. Because the administrative unit represented by each row cannot be inferred safely from a column label alone, that unit must be confirmed from the source documentation before submission.\n\n"
            f"The declared outcome was `{protocol.get('outcome') or 'not selected'}` and the declared predictors were {protocol.get('predictors') or 'not selected'}. The analytical equation was `{protocol.get('equation') or 'not specified'}`. Variables were retained in their recorded units unless the protocol documented a safe derived expression. Descriptive statistics and variable-level missingness were calculated for the complete selected scope; exploratory monotonic relationships were assessed using Spearman correlations and false-discovery control where a headline association was selected from multiple pairs.\n\n"
            f"The primary executed routine was {current_result.algorithm}. Model-based results, where present, were accompanied by sample size, uncertainty and fit evidence. Regression coefficients were interpreted as conditional associations unless the research design separately established temporal ordering and a credible identification strategy. Robustness assessment was organised around missing-data sensitivity, alternative operationalisations, influential observations, covariance assumptions and reproducible simulation settings.\n\n"
            "All outputs were generated from the archived filtered dataset and saved protocol. The bundle preserves the selected PDF pages, transformations, equations, seeds, repetitions, result tables, figure data and software-facing settings needed to reconstruct the analysis."
        )
        reply += "\n\nThe downloadable Methodology outline table remains as the subsection-level audit trail behind this draft."
        return result, reply

    if results_outline_intent:
        guide = build_output_guide(current_result.tables)
        result = _result_copy(current_result, "Results-section outline")
        result.tables["Results output guide"] = guide
        profile = scope_profile(df).iloc[0]
        ranked, _ = rank_striking_findings(df, protocol, current_result)
        missing = _missingness_table(df)
        reply = (
            "**Draft results text and section order**\n\n"
            f"**1. Analytical sample and data quality.** The retained dataset contained {int(profile.records):,} observations and {int(profile.variables):,} variables, of which {int(profile.numeric_variables):,} were numeric. Missing values accounted for {float(profile.missing_percent):.2f}% of all cells and {int(profile.duplicate_rows):,} exact duplicate rows were detected."
        )
        if not missing.empty and float(missing.iloc[0].missing_percent) > 0:
            worst = missing.iloc[0]
            reply += f" The largest variable-specific deficit concerned `{worst.variable}`, which was {float(worst.missing_percent):.2f}% missing ({int(worst.observed):,} observed values)."
        if not ranked.empty:
            top = ranked.iloc[0]
            reply += (
                f"\n\n**2. Principal exploratory finding.** The strongest adequately covered non-identifier association was observed between `{top.variable_1}` and `{top.variable_2}` (Spearman ρ={float(top.estimate):.3f}, n={int(top.n):,}, {_p_text(float(top.p_value))}, {_q_text(float(top.q_value_bh))}). The magnitude indicates substantial common movement and possible construct overlap; it was therefore treated as associative evidence rather than an independent causal effect."
            )
        trend = current_result.tables.get("Longitudinal results", pd.DataFrame())
        if not trend.empty and "year" in trend:
            years = pd.to_numeric(trend.year, errors="coerce").dropna()
            reply += f"\n\n**3. Longitudinal evidence.** The saved trend table covered {int(years.min())}–{int(years.max())} across {years.nunique()} observed years. Changes were reported as temporal movements and not attributed to time itself or to an intervention without a separate identification design."
        fit = current_result.tables.get("OLS fit", pd.DataFrame())
        if not fit.empty:
            reply += f"\n\n**4. Conditional model.** The HC3 OLS model used {int(fit.iloc[0].n):,} complete observations and achieved adjusted R²={float(fit.iloc[0].adjusted_r_squared):.3f}. Coefficients and intervals should be presented next, followed by specification and influence diagnostics."
        reply += "\n\n**5. Robustness and boundary of inference.** Any sensitivity, simulation or alternative specification should follow the primary result. Statistical significance was not treated as proof of practical importance, generalisability or causality."
        reply += "\n\nThe Results output guide table gives the matching purpose, reading rule, non-claim and next step for every generated table."
        return result, reply

    if pdf_methodology_intent:
        methods = _method_evidence_table(evidence)
        result = _result_copy(current_result, "PDF methodology extraction and replication audit")
        if not methods.empty:
            methods = methods.copy()
            methods["replication_status"] = methods.method.map(lambda value: "Supported in a specialist dashboard module" if value in ["OLS / linear regression", "Panel fixed effects", "Monte Carlo / simulation", "Clustering", "Spatial analysis", "MCDA", "Difference-in-differences", "Instrumental variables"] else "Requires manual implementation review")
            result.tables["PDF methodology extraction"] = methods
        if evidence.empty:
            return result, "No PDF pages are saved in the current evidence scope, so no paper methodology can be extracted or replicated."
        detected = ", ".join(methods.method.drop_duplicates()) if not methods.empty else "no named supported estimator"
        reply = (
            f"**Methodology extraction result**\n\nThe selected PDF evidence contains {len(evidence):,} page(s). Detected methods: {detected}. "
            "The new extraction table records document, page, snippet and replication status.\n\n"
            "**Replication rule:** the app may reproduce an implemented estimator only after the paper's outcome, predictors, transformations, sample restrictions, covariance rule and diagnostics are mapped to actual spreadsheet columns. Missing elements are reported as deviations; they are never invented. "
            "A textual method mention alone is insufficient for numerical replication."
        )
        return result, reply

    if equation_audit_intent:
        equation = protocol.get("equation", "").strip()
        expression = protocol.get("derived_expression", "").strip()
        rows = [
            {"audit_item": "Equation supplied", "status": "Present" if equation else "Missing", "explanation": equation or "No equation is saved."},
            {"audit_item": "Outcome mapped", "status": "Mapped" if outcome else "Missing", "explanation": outcome or "Select a substantive outcome."},
            {"audit_item": "Predictors mapped", "status": "Mapped" if predictors else "Missing", "explanation": ", ".join(predictors) if predictors else "Select one or more predictors."},
            {"audit_item": "Computable expression", "status": "Present" if expression else "Not requested", "explanation": expression or "The LaTeX equation is documentation unless a safe derived expression is supplied."},
            {"audit_item": "Causal identification", "status": "Not established by equation alone", "explanation": "An equation specifies a model; causal interpretation additionally requires design, timing and identification assumptions."},
            {"audit_item": "Arbitrary algorithm execution", "status": "Blocked unless implemented", "explanation": "Free-form prose or code is documented and checked, but only validated dashboard estimators are executed."},
        ]
        audit = pd.DataFrame(rows)
        result = _result_copy(current_result, "Equation and algorithm audit")
        result.tables["Equation and algorithm audit"] = audit
        reply = "**Equation/algorithm audit**\n\n" + "\n".join(f"- **{row.audit_item}: {row.status}.** {row.explanation}" for row in audit.itertuples())
        reply += "\n\nThis is a structural correctness audit. A mathematical expression can be computed only through the restricted derived-variable engine; unsupported estimators are not simulated or represented as validated."
        return result, reply

    if normality_intent:
        from analytics_core import normality_tests
        variables = mentioned_numeric or predictors or numeric_all[:20]
        table = normality_tests(df, variables)
        result = _result_copy(current_result, "Normality diagnostics")
        result.tables["Normality tests"] = table
        rejected = table[(table.test != "Anderson–Darling") & (table.p_value < .05)] if not table.empty else table
        reply = (
            f"**Feasible and executed.** Normality diagnostics were run for {len(variables):,} numeric variable(s). "
            f"Among tests with p-values, {len(rejected):,} result(s) reject normality at the 5% level. "
            "This does not automatically invalidate regression—large-sample coefficient inference concerns residual/model behaviour, not whether every raw variable is Gaussian. "
            "Use distribution plots, robust inference and transformations only where their scientific meaning is defensible."
        )
        return result, reply

    if outlier_intent:
        from analytics_core import outlier_summary
        variables = mentioned_numeric or predictors or numeric_all[:30]
        table = outlier_summary(df, variables)
        result = _result_copy(current_result, "Outlier and extreme-value audit")
        result.tables["Outlier audit"] = table
        reply = (
            f"**Feasible and executed.** The outlier audit covers {len(variables):,} numeric variable(s). "
            "Flagged values are diagnostic candidates, not automatic errors: retain, correct or exclude them only after checking provenance, units and influence on the estimand. "
            "The full variable-level counts and thresholds are included in the results bundle."
        )
        return result, reply

    if group_test_intent:
        if not outcome or not group_column:
            return current_result, "**Not feasible — logic/setup error.** A group comparison requires a numeric outcome and a grouping variable; the Chair will not guess either role."
        from analytics_core import group_tests
        table = group_tests(df, [outcome], group_column)
        result = _result_copy(current_result, "Group-comparison tests")
        result.tables["Group comparison tests"] = table
        if table.empty:
            return result, "**Not feasible with the retained evidence.** Fewer than two analysable groups with adequate observations remain."
        primary = table.iloc[0]
        reply = (
            f"**Feasible and executed.** `{outcome}` was compared across `{group_column}`. "
            f"The leading {primary.test} result is statistic={float(primary.statistic):.3g}, p={float(primary.p_value):.3g}; "
            f"the reported {primary.effect_metric} is {float(primary.effect_size):.3g} where estimable. "
            "The effect size describes the magnitude of observed group separation; it does not show that group membership caused the difference."
        )
        return result, reply

    if pca_intent:
        from analytics_core import pca_table
        variables = mentioned_numeric or predictors or [c for c in numeric_all if not _identifier_or_time_like(c, df[c])][:30]
        if len(variables) < 2:
            return current_result, "**Not feasible — logic/setup error.** PCA requires at least two varying numeric indicators."
        loadings, variance = pca_table(df, variables, n_components=min(8, len(variables)))
        result = _result_copy(current_result, "Principal component analysis")
        result.tables["PCA loadings"] = loadings
        result.tables["PCA explained variance"] = variance
        first = float(variance.iloc[0].explained_variance_ratio) if not variance.empty else np.nan
        reply = (
            f"**Feasible and executed.** PCA used {len(variables):,} median-imputed, standardised indicators. "
            f"The first component explains {100 * first:.1f}% of total standardised variance. "
            "The loading table shows which variables define each component; a component should receive a substantive label only when its high-loading variables form a coherent construct. "
            "PCA summarises covariance and does not prove an underlying causal mechanism."
        )
        return result, reply

    if reliability_intent:
        from analytics_core import cronbach_alpha
        items = mentioned_numeric or predictors
        if len(items) < 2:
            return current_result, "**Not feasible — logic/setup error.** Reliability analysis requires at least two conceptually related items selected as predictors."
        summary, item_table = cronbach_alpha(df, items)
        result = _result_copy(current_result, "Scale reliability analysis")
        result.tables["Cronbach alpha summary"] = summary
        result.tables["Cronbach item diagnostics"] = item_table
        alpha = float(summary.iloc[0].cronbach_alpha)
        reply = (
            f"**Feasible and executed.** Cronbach’s α = {alpha:.3f} across {int(summary.iloc[0].items)} item(s) and "
            f"{int(summary.iloc[0].complete_cases):,} complete cases. "
            "Alpha measures internal consistency, not unidimensionality or validity. Inspect item–total correlations, alpha-if-deleted and a theory/PCA assessment before combining the items."
        )
        return result, reply

    if clustering_intent:
        from advanced_analytics import advanced_clustering
        variables = mentioned_numeric or predictors or [c for c in numeric_all if not _identifier_or_time_like(c, df[c])][:12]
        if not variables:
            return current_result, "**Not feasible — logic/setup error.** Clustering requires at least one varying numeric indicator."
        seed_match = re.search(r"seed\s*[:=]?\s*(\d+)", q)
        seed = int(seed_match.group(1)) if seed_match else 42
        output = advanced_clustering(df, variables, automatic_k=True, max_k=8, seed=seed)
        result = _result_copy(current_result, "Automatic K-means clustering")
        result.tables["Cluster assignments"] = output.assignments
        result.tables["Cluster profiles"] = output.profiles
        result.tables["Cluster diagnostics"] = output.diagnostics
        result.tables["Cluster embedding"] = output.embedding
        reply = "**Feasible and executed.** " + " ".join(output.interpretation)
        return result, reply

    if prediction_intent:
        if not outcome or not predictors:
            return current_result, "**Not feasible — logic/setup error.** Predictive comparison requires an explicit continuous outcome and non-leaking predictors."
        from advanced_analytics import predictive_model_comparison
        performance, predictions_table, importance, comments = predictive_model_comparison(df, outcome, predictors, folds=5, seed=42)
        result = _result_copy(current_result, "Cross-validated predictive comparison")
        result.tables["Predictive performance"] = performance
        result.tables["Cross-validated predictions"] = predictions_table
        result.tables["Permutation importance"] = importance
        reply = "**Feasible and executed.** " + " ".join(comments)
        return result, reply

    if time_series_intent and not forecast_intent:
        if not year_column:
            return current_result, "**Not feasible — logic/setup error.** Time-series diagnostics require a mapped time variable."
        from analytics_core import time_series_tests
        variables = mentioned_numeric or ([outcome] if outcome else []) or predictors or [c for c in numeric_all if c != year_column][:12]
        table = time_series_tests(df, variables, year_column)
        result = _result_copy(current_result, "Time-series stationarity diagnostics")
        result.tables["Time-series tests"] = table
        if table.empty:
            return result, "**Not feasible with the retained evidence.** No selected series has the minimum ordered observations and variation required by the implemented ADF/KPSS diagnostics."
        reply = (
            f"**Feasible and executed.** ADF and KPSS diagnostics were attempted for {len(variables):,} selected series ordered by `{year_column}`. "
            "ADF tests the null of a unit root, whereas KPSS tests the null of stationarity; the two should be read together because they reverse the null hypothesis. "
            "The complete statistics, lags, sample sizes and p-values are provided in the time-series table. Aggregated annual data with very few years cannot support a credible forecast merely because many project rows exist."
        )
        return result, reply

    if forecast_intent:
        if not year_column or not outcome:
            return current_result, "**Not feasible — logic/setup error.** ARIMA forecasting requires a mapped time variable and numeric outcome."
        from analytics_core import arima_forecast
        horizon_match = re.search(r"(\d+)\s*(?:periods|steps|years|months)", q)
        steps = int(np.clip(int(horizon_match.group(1)), 1, 50)) if horizon_match else 5
        forecast, fit = arima_forecast(df, outcome, year_column, (1, 1, 1), steps=steps)
        result = _result_copy(current_result, "ARIMA(1,1,1) forecast")
        result.tables["ARIMA forecast"] = forecast
        result.tables["ARIMA fit"] = fit
        reply = (
            f"**Feasible and executed.** An ARIMA(1,1,1) model generated {steps} forecast step(s) for `{outcome}`. "
            f"Model AIC={float(fit.iloc[0].aic):.3f} and BIC={float(fit.iloc[0].bic):.3f}. "
            "Forecast intervals represent model-based uncertainty conditional on the ordering, specification and absence of an unmodelled structural break; they are not policy-effect estimates."
        )
        return result, reply

    if panel_intent:
        if not (group_column and year_column and outcome and predictors):
            return current_result, "**Not feasible — logic/setup error.** Panel estimation requires entity/group, time, outcome and predictor mappings."
        from advanced_analytics import panel_model_suite
        fit, coefficients, hausman, panel_data, comments = panel_model_suite(
            df, group_column, year_column, outcome, predictors, aggregation="Mean", covariance="Clustered by entity"
        )
        result = _result_copy(current_result, "Panel model suite")
        result.tables["Panel coefficients"] = coefficients
        result.tables["Panel model fit"] = fit
        result.tables["Panel Hausman test"] = hausman
        result.tables["Panel analytical data"] = panel_data
        reply = "**Feasible and executed.** " + " ".join(comments)
        return result, reply

    if monte_carlo_intent:
        if not outcome or not predictors:
            result = _result_copy(current_result, "Monte Carlo request — variables required")
            return result, (
                "A Monte Carlo OLS simulation was not executed because the saved protocol does not contain both an explicit outcome and at least one predictor. "
                "Select a substantively meaningful Outcome / principal measure and Predictor set, press Run Research Command, then repeat the request. "
                "The app will not guess a dependent variable."
            )
        repetitions_match = re.search(r"([\d,]+)\s*(?:replications|repetitions|simulations|draws|επαναλήψεις)", q)
        seed_match = re.search(r"seed\s*[:=]?\s*(\d+)", q)
        simulations = int(repetitions_match.group(1).replace(",", "")) if repetitions_match else 2_000
        simulations = int(np.clip(simulations, 100, 20_000))
        seed = int(seed_match.group(1)) if seed_match else 42
        method = "Residual bootstrap" if "residual" in q else "Parametric normal" if "parametric" in q or "normal simulation" in q else "Wild bootstrap"
        from analytics_core import monte_carlo_ols
        summary, draws, fit = monte_carlo_ols(
            df, outcome, predictors, simulations=simulations, method=method, seed=seed,
        )
        result = _result_copy(current_result, f"Monte Carlo OLS — {method}")
        result.tables["Monte Carlo summary"] = summary
        result.tables["Monte Carlo draws"] = draws
        result.tables["Monte Carlo fit and settings"] = fit
        result.comments.append(
            f"Monte Carlo OLS used {simulations:,} reproducible {method.lower()} draws with seed {seed}."
        )
        result.comments.append(
            "Simulation measures uncertainty conditional on the selected model and resampling mechanism; it does not correct confounding, measurement error or an invalid outcome/predictor choice."
        )
        interval_columns = [c for c in summary if c.startswith("ci_")]
        low_column = next((c for c in interval_columns if c.endswith("_low")), None)
        high_column = next((c for c in interval_columns if c.endswith("_high")), None)
        statements = []
        for _, row in summary[summary.term != "const"].head(6).iterrows():
            low = row[low_column] if low_column else np.nan
            high = row[high_column] if high_column else np.nan
            statements.append(
                f"`{row['term']}`: OLS estimate {float(row['ols_estimate']):.4g}; simulated 95% interval "
                f"[{float(low):.4g}, {float(high):.4g}]; probability positive {100 * float(row['probability_positive']):.1f}%"
            )
        reply = (
            f"**Monte Carlo analysis executed**\n\nMethod: {method}; replications: {simulations:,}; seed: {seed}; "
            f"complete observations: {int(fit.iloc[0].complete_observations):,}.\n\n"
            + "\n\n".join(statements)
            + "\n\n**How to interpret it.** An interval entirely above zero indicates a stable positive sign under this simulation design; an interval entirely below zero indicates a stable inverse sign; an interval crossing zero indicates that the sign is not stable. "
              "These are uncertainty results for the stated OLS model, not proof of causality. The full summary, every simulation draw and all settings are now included in the XLSX and Research Chair bundle."
        )
        return result, reply

    if regression_intent and not paper_intent:
        if not outcome or not predictors:
            result = execute_protocol(df, "Descriptive profile", equation=equation)
            reply = _command_reply("pre-regression data audit", protocol, result, evidence)
            reply += "\n\nOLS was not estimated because choosing the dependent variable is a scientific decision. Select Outcome / principal measure and at least one Predictor, press Run Research Command, and ask again."
            return result, reply
        result = execute_protocol(df, "OLS specification", outcome, predictors, year_column, group_column, aggregation, equation)
        return result, _command_reply("HC3 OLS specification", protocol, result, evidence)

    if trend_intent and not paper_intent:
        if not year_column:
            result = execute_protocol(df, "Descriptive profile", equation=equation)
            reply = _command_reply("descriptive audit", protocol, result, evidence)
            reply += "\n\nA longitudinal table was not generated because no valid year/date variable is saved. Select the year variable and rerun the command."
            return result, reply
        measures = mentioned_numeric or list(dict.fromkeys(([outcome] if outcome else []) + predictors)) or numeric_all
        measures = [c for c in measures if c != year_column][:12]
        result = execute_protocol(df, "Longitudinal trend", measures[0] if measures else None, measures[1:], year_column, group_column, aggregation, equation)
        return result, _command_reply("longitudinal trend analysis", protocol, result, evidence)

    if correlation_intent and not paper_intent:
        measures = mentioned_numeric or list(dict.fromkeys(([outcome] if outcome else []) + predictors)) or numeric_all[:25]
        if len(measures) < 2:
            result = execute_protocol(df, "Descriptive profile", equation=equation)
            return result, _command_reply("descriptive audit", protocol, result, evidence) + "\n\nAt least two numeric variables are required for correlation screening."
        result = execute_protocol(df, "Correlation screening", measures[0], measures[1:], year_column, group_column, aggregation, equation)
        table = result.tables.get("Correlation screening", pd.DataFrame())
        if table.empty:
            return result, "**Feasible but no estimable pair remained.** At least two selected numeric variables need three paired, varying observations."
        count = len(table)
        order = np.argsort(table.p_value.to_numpy(float))
        adjusted_sorted = np.minimum.accumulate((table.p_value.to_numpy(float)[order] * count / np.arange(1, count + 1))[::-1])[::-1]
        adjusted = np.empty_like(adjusted_sorted); adjusted[order] = np.clip(adjusted_sorted, 0, 1)
        table = table.copy(); table["q_value_bh"] = adjusted
        result.tables["Correlation screening"] = table
        positive = table.sort_values("spearman_rho", ascending=False).iloc[0]
        inverse_candidates = table[table.spearman_rho < 0]
        inverse = inverse_candidates.sort_values("spearman_rho").iloc[0] if not inverse_candidates.empty else None
        weakest = table.iloc[(table.spearman_rho.abs()).argmin()]
        pieces = [
            f"**Strongest positive relationship:** `{positive.variable_1}` with `{positive.variable_2}` (ρ={float(positive.spearman_rho):.3f}, n={int(positive.n):,}, {_p_text(float(positive.p_value))}, {_q_text(float(positive.q_value_bh))}). Both measures tend to rise together.",
        ]
        if inverse is not None:
            pieces.append(f"**Strongest inverse relationship:** `{inverse.variable_1}` with `{inverse.variable_2}` (ρ={float(inverse.spearman_rho):.3f}, n={int(inverse.n):,}, {_p_text(float(inverse.p_value))}, {_q_text(float(inverse.q_value_bh))}). Higher values of one tend to accompany lower values of the other.")
        else:
            pieces.append("**Strongest inverse relationship:** none of the screened pairs had a negative Spearman coefficient.")
        pieces.append(f"**Weakest relationship:** `{weakest.variable_1}` with `{weakest.variable_2}` (ρ={float(weakest.spearman_rho):.3f}, n={int(weakest.n):,}). A value close to zero indicates little monotonic association, although a nonlinear relationship may still exist.")
        pieces.append("Correlation measures ordered co-movement. It does not show which variable comes first, rule out confounding or establish a causal mechanism.")
        return result, "\n\n".join(pieces)

    if quality_intent and not paper_intent:
        result = execute_protocol(df, "Descriptive profile", equation=equation)
        missing = result.tables.get("Variable missingness", pd.DataFrame()).copy()
        if not missing.empty:
            missing["modelling_recommendation"] = pd.cut(
                missing.missing_percent,
                bins=[-0.001, 0, 5, 20, 50, 100.001],
                labels=[
                    "Complete: ordinarily usable after validity checks",
                    "Low missingness: usable with transparent complete-case count",
                    "Material missingness: use sensitivity analysis",
                    "Severe missingness: secondary/sensitivity use only without a defensible strategy",
                    "Extreme missingness: exclude from primary modelling unless values can be recovered",
                ],
                include_lowest=True,
            ).astype(str)
            result.tables["Variable missingness and usability"] = missing
            affected = missing[missing.missing_percent > 0].head(8)
            if affected.empty:
                detail = "No retained variable contains a missing value. Completeness does not, however, establish measurement validity."
            else:
                detail = "\n".join(
                    f"- `{row.variable}`: {float(row.missing_percent):.2f}% missing ({int(row.observed):,} observed) — {row.modelling_recommendation}."
                    for row in affected.itertuples(index=False)
                )
        else:
            detail = "No variable-level missingness table could be produced."
        reply = (
            "**Data-quality answer**\n\n"
            + detail
            + "\n\nThese thresholds are transparent screening rules, not automatic deletion instructions. A variable with low missingness can still be invalid, and a sparse variable may remain scientifically valuable if the missing values can be recovered or modelled under a defensible mechanism. The bundle includes the complete usability classification for every retained variable."
        )
        return result, reply

    if pdf_intent and not paper_intent:
        result = _result_copy(current_result, "PDF method evidence review")
        methods = _method_evidence_table(evidence)
        if not methods.empty:
            result.tables["Methods found in selected PDFs"] = methods
        reply = _command_reply("selected-PDF method review", protocol, result, evidence)
        if evidence.empty:
            reply += "\n\nNo PDF pages are saved in the current evidence scope."
        elif methods.empty:
            reply += "\n\nThe selected pages were searched, but no supported method label was detected. Inspect the downloadable page evidence for terminology not covered by the deterministic dictionary."
        return result, reply

    # Paper/full-analysis commands and otherwise-unspecified questions receive
    # a concrete baseline analysis rather than generic writing advice.
    result = execute_protocol(df, "Descriptive profile", equation=equation)
    result.algorithm = "Comprehensive paper-ready baseline"
    correlation_variables = mentioned_numeric or list(dict.fromkeys(([outcome] if outcome else []) + predictors)) or numeric_all[:25]
    if len(correlation_variables) >= 2:
        correlations, total = _correlation_table(df, correlation_variables)
        if not correlations.empty:
            result.tables["Correlation screening"] = correlations
            result.comments.append("Spearman screening is exploratory and does not establish causation.")
        if total > 25:
            result.comments.append(f"Correlation screening was bounded to the first 25 of {total} candidate variables.")
    if year_column:
        measures = [c for c in (mentioned_numeric or list(dict.fromkeys(([outcome] if outcome else []) + predictors)) or numeric_all) if c != year_column][:12]
        trend = _longitudinal_table(df, measures, year_column, group_column, aggregation)
        if not trend.empty:
            result.tables["Longitudinal results"] = trend
            result.comments.append("Longitudinal movement is descriptive without a defensible identification strategy.")
    if outcome and predictors:
        ols = execute_protocol(df, "OLS specification", outcome, predictors, year_column, group_column, aggregation, equation)
        for name in ["OLS coefficients", "OLS fit"]:
            if name in ols.tables:
                result.tables[name] = ols.tables[name]
        result.comments.extend(comment for comment in ols.comments if comment not in result.comments)
    methods = _method_evidence_table(evidence)
    if not methods.empty:
        result.tables["Methods found in selected PDFs"] = methods
    action = "comprehensive paper-ready baseline" if paper_intent else "evidence-grounded exploratory analysis"
    return result, _command_reply(action, protocol, result, evidence)


def ollama_models(endpoint: str = "http://127.0.0.1:11434", timeout: float = .35) -> list[str]:
    import requests
    try:
        response = requests.get(f"{endpoint.rstrip('/')}/api/tags", timeout=timeout)
        response.raise_for_status()
        return [item["name"] for item in response.json().get("models", []) if item.get("name")]
    except Exception:
        return []


def ollama_reply(
    question: str,
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
    model: str,
    endpoint: str = "http://127.0.0.1:11434",
    timeout: int = 120,
) -> str:
    import requests
    context_tables = "\n\n".join(f"TABLE: {name}\n{_compact_table(table)}" for name, table in result.tables.items())
    prompt = f"""You are an academic research-methods assistant. Use only the supplied context. Do not invent results or citations. Distinguish association, prediction and causality. Write equations in LaTeX delimiters. State limitations explicitly. Answer in the language of the question.

PROTOCOL
{json.dumps(protocol, ensure_ascii=False, default=str)}

RESULTS
{context_tables[:18000]}

SELECTED PDF EVIDENCE
{evidence_text(evidence, 14000)}

QUESTION
{question}
"""
    response = requests.post(
        f"{endpoint.rstrip('/')}/api/generate",
        json={"model": model, "prompt": prompt, "stream": False, "options": {"temperature": .15}},
        timeout=timeout,
    )
    response.raise_for_status()
    return str(response.json().get("response", "")).strip()


def dataframe_markdown(table: pd.DataFrame, max_rows: int = 20) -> str:
    if table is None or table.empty:
        return "_No results generated._"
    show = table.head(max_rows).copy()
    columns = [str(c) for c in show.columns]
    lines = ["| " + " | ".join(columns) + " |", "| " + " | ".join(["---"] * len(columns)) + " |"]
    for row in show.itertuples(index=False, name=None):
        cells = [str(value).replace("|", "\\|").replace("\n", " ") for value in row]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def build_paper_blueprint(
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
    reply: str = "",
) -> str:
    question = protocol.get("research_question") or "[Insert one principal research question]"
    title = protocol.get("working_title") or "Research paper blueprint"
    steps = protocol.get("steps") or "Document preprocessing, estimation, diagnostics and robustness in execution order."
    limitations = protocol.get("limitations") or "State sampling, measurement, missing-data, specification and causal-identification limitations."
    tables = []
    for name, table in result.tables.items():
        tables.append(f"### {name}\n\n{dataframe_markdown(table)}")
    sources = "\n".join(f"- {row.document}, p. {row.page}" for row in evidence.drop_duplicates(["document", "page"]).itertuples()) or "- No PDF pages selected."
    guide = build_output_guide(
        result.tables,
        ["Selected longitudinal evidence chart"] if not result.tables.get("Longitudinal results", pd.DataFrame()).empty else [],
    )
    guide_text = output_guide_markdown(guide).replace("# Output interpretation guide", "## How to read every table and chart", 1)
    return f"""# {title}

## Proposed contribution

Use the selected primary data and documentary evidence to answer the stated question with a transparent, reproducible protocol. Originality must arise from the question, data, identification and comparison—not from the software alone.

## Research question

{question}

## Hypotheses and theoretical expectations

Translate the research question into directional or non-directional hypotheses before interpreting significance. Connect each hypothesised mechanism to an observable variable and an expected sign or pattern.

## Data and analytical scope

{dataframe_markdown(result.tables.get('Scope profile', pd.DataFrame()))}

Explain the unit of observation, inclusion and exclusion rules, time coverage, transformations, missing-data handling and any joins. The exported filtered dataset is the exact analytical scope.

## Algorithm and reproducible steps

**Algorithm:** {result.algorithm}

{steps}

## Equation and operationalisation

$$
{result.equation or r'Y_i = \beta_0 + \sum_{k=1}^{K}\beta_k X_{ki} + \varepsilon_i'}
$$

Executed safe expression: `{result.executed_expression or 'None'}`

Define every symbol, unit, transformation, reference category, aggregation and expected coefficient interpretation.

## Results to report

{chr(10).join(tables)}

Report point estimates together with uncertainty, sample size, diagnostics and substantive magnitude. Do not select models solely because they return smaller p-values.

## Interpretation

{reply or 'Interpret the generated tables in relation to the research question, theory and units of measurement.'}

{guide_text}

## Robustness and validation plan

1. Re-estimate the primary specification under defensible missing-data and covariance rules.
2. Test sensitivity to influential observations, alternative operationalisations and temporal windows.
3. Compare model families only where they answer the same estimand.
4. Preserve seeds, software versions, selected pages, filters and the exported analytical dataset.
5. Distinguish exploratory, confirmatory, predictive and causal claims.

## Limitations

{limitations}

## Recommended paper structure

1. Introduction: problem, gap, question, contribution and principal finding.
2. Literature and theory: mechanisms and hypotheses, not a catalogue of sources.
3. Data and methods: provenance, sample construction, variables, equation, assumptions and reproducibility.
4. Results: descriptive evidence, primary model, diagnostics and robustness.
5. Discussion: theoretical and policy meaning, comparisons and boundary conditions.
6. Conclusion: answer, contribution, limitations and next research step.

## Selected documentary evidence

{sources}

Selected PDF passages are research notes, not automatically verified quotations. Check exact wording and bibliographic metadata in the original documents before submission.
"""


def docx_bytes(markdown: str) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Word report export requires python-docx. Install the bundled requirements.") from exc
    document = Document()
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            document.add_paragraph()
        elif line.startswith("# "):
            document.add_heading(line[2:], 0)
        elif line.startswith("## "):
            document.add_heading(line[3:], 1)
        elif line.startswith("### "):
            document.add_heading(line[4:], 2)
        elif re.match(r"^\d+\. ", line):
            document.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|") or line.startswith("$$"):
            document.add_paragraph(line)
        else:
            document.add_paragraph(line.replace("**", ""))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _chair_figure_specs(scoped_data: pd.DataFrame, result: ProtocolResult) -> list[dict[str, Any]]:
    """Create a bounded set of chart specifications from computed Chair tables."""
    specs: list[dict[str, Any]] = []
    missing = next((table for name, table in result.tables.items() if "missingness" in name.casefold() and not table.empty), pd.DataFrame())
    if not missing.empty and {"variable", "missing_percent"}.issubset(missing.columns):
        affected = missing[pd.to_numeric(missing["missing_percent"], errors="coerce") > 0].copy()
        data = affected.nlargest(min(20, len(affected)), "missing_percent").sort_values("missing_percent")
    else:
        data = pd.DataFrame()
    if not data.empty:
        top = data.iloc[-1]
        specs.append({
            "stem": "missingness_profile", "kind": "barh", "data": data,
            "x": "missing_percent", "y": "variable", "title": "Variables with the greatest missingness",
            "x_label": "Missing observations (%)",
            "commentary": f"The largest missing-data burden is `{top.variable}` at {float(top.missing_percent):.2f}%. This chart identifies where effective sample size and selection bias may become material; it does not establish why values are absent.",
        })

    correlations = next((table for name, table in result.tables.items() if ("correlation" in name.casefold() or "ranked statistical" in name.casefold()) and not table.empty), pd.DataFrame())
    if not correlations.empty:
        if {"variable_1", "variable_2", "spearman_rho"}.issubset(correlations.columns):
            data = correlations.head(20).copy()
            data["pair"] = data.variable_1.astype(str) + " ↔ " + data.variable_2.astype(str)
            data = data.sort_values("spearman_rho")
            specs.append({
                "stem": "strongest_associations", "kind": "signed_barh", "data": data,
                "x": "spearman_rho", "y": "pair", "title": "Strongest screened monotonic associations",
                "x_label": "Spearman correlation (ρ)",
                "commentary": "Positive bars indicate variables moving in the same direction; negative bars indicate inverse movement; values near zero indicate little monotonic relationship. Correlation does not establish direction, mechanism or causality.",
            })
        elif {"variable_1", "variable_2", "estimate"}.issubset(correlations.columns):
            data = correlations.head(20).copy()
            data["pair"] = data.variable_1.astype(str) + " ↔ " + data.variable_2.astype(str)
            data = data.sort_values("estimate")
            specs.append({
                "stem": "ranked_statistical_findings", "kind": "signed_barh", "data": data,
                "x": "estimate", "y": "pair", "title": "Ranked statistical findings",
                "x_label": "Spearman correlation (ρ)",
                "commentary": "The figure ranks adequately covered, non-identifier pairs by association magnitude. Very large values may indicate construct overlap or mechanically related indicators rather than an independent scientific effect.",
            })

    trend = next((table for name, table in result.tables.items() if "longitudinal" in name.casefold() and not table.empty), pd.DataFrame())
    if not trend.empty and "year" in trend:
        numeric = [c for c in trend.select_dtypes(include=np.number).columns if c != "year"][:8]
        if numeric:
            identifiers = [c for c in trend.columns if c not in numeric and c != "year"]
            data = trend.melt(id_vars=["year", *identifiers], value_vars=numeric, var_name="measure", value_name="value")
            specs.append({
                "stem": "longitudinal_evidence", "kind": "line", "data": data,
                "x": "year", "y": "value", "group": "measure", "title": "Selected longitudinal evidence",
                "x_label": "Year", "y_label": "Aggregated value",
                "commentary": "Each line shows the saved aggregation across observed years. A rising or falling line is temporal description; it does not identify the cause of change and may reflect sample composition, reporting or programme design.",
            })

    coefficients = next((table for name, table in result.tables.items() if ("ols coefficients" in name.casefold() or "panel coefficients" in name.casefold()) and not table.empty), pd.DataFrame())
    if not coefficients.empty and {"term", "coefficient"}.issubset(coefficients.columns):
        data = coefficients[coefficients.term.astype(str) != "const"].head(30).copy().sort_values("coefficient")
        if not data.empty:
            specs.append({
                "stem": "coefficient_estimates", "kind": "coefficient", "data": data,
                "x": "coefficient", "y": "term", "title": "Coefficient estimates and 95% intervals",
                "x_label": "Estimated coefficient",
                "commentary": "Points are conditional coefficient estimates and horizontal lines are 95% intervals where available. Intervals crossing zero indicate that the sign is not statistically resolved at the conventional 5% level; neither sign nor significance establishes causality.",
            })

    draws = next((table for name, table in result.tables.items() if "monte carlo draws" in name.casefold() and not table.empty), pd.DataFrame())
    if not draws.empty:
        terms = [c for c in draws.select_dtypes(include=np.number).columns if c not in ["simulation", "const"]]
        if terms:
            term = terms[0]
            data = draws[[term]].rename(columns={term: "coefficient"}).dropna()
            specs.append({
                "stem": "monte_carlo_distribution", "kind": "hist", "data": data,
                "x": "coefficient", "title": f"Monte Carlo coefficient distribution: {term}",
                "x_label": "Simulated coefficient", "y_label": "Frequency",
                "commentary": "The distribution shows how the selected coefficient varies under the declared resampling mechanism. Concentration on one side of zero indicates sign stability within that simulation design, not freedom from confounding or measurement error.",
            })

    variance = next((table for name, table in result.tables.items() if "pca explained variance" in name.casefold() and not table.empty), pd.DataFrame())
    if not variance.empty and {"component", "explained_variance_ratio"}.issubset(variance.columns):
        specs.append({
            "stem": "pca_explained_variance", "kind": "bar", "data": variance.copy(),
            "x": "component", "y": "explained_variance_ratio", "title": "PCA explained variance",
            "x_label": "Principal component", "y_label": "Share of standardised variance",
            "commentary": "Bars show the share of standardised variance captured by each component. Retaining a component is a compression decision, not evidence that the component is a valid latent construct.",
        })

    outliers = next((table for name, table in result.tables.items() if "outlier audit" in name.casefold() and not table.empty), pd.DataFrame())
    if not outliers.empty and {"variable", "outlier_pct"}.issubset(outliers.columns):
        data = outliers.head(20).sort_values("outlier_pct")
        specs.append({
            "stem": "outlier_profile", "kind": "barh", "data": data,
            "x": "outlier_pct", "y": "variable", "title": "Robust univariate outlier flags",
            "x_label": "IQR-flagged observations (%)",
            "commentary": "Bars show the share of observations beyond the conventional 1.5×IQR fences. A flagged observation may be genuine, influential or erroneous; the chart is a provenance-checking device and not a deletion rule.",
        })

    performance = next((table for name, table in result.tables.items() if "predictive performance" in name.casefold() and not table.empty), pd.DataFrame())
    if not performance.empty and {"model", "cross_validated_rmse"}.issubset(performance.columns):
        data = performance.sort_values("cross_validated_rmse", ascending=False)
        specs.append({
            "stem": "predictive_performance", "kind": "barh", "data": data,
            "x": "cross_validated_rmse", "y": "model", "title": "Cross-validated predictive error",
            "x_label": "Cross-validated RMSE (lower is better)",
            "commentary": "Shorter bars indicate lower out-of-sample prediction error under the saved folds and predictor set. Better prediction does not turn a variable into a causal determinant and may be optimistic when rows are dependent by entity or time.",
        })

    embedding = next((table for name, table in result.tables.items() if "cluster embedding" in name.casefold() and not table.empty), pd.DataFrame())
    if not embedding.empty and {"dimension_1", "dimension_2", "cluster_label"}.issubset(embedding.columns):
        specs.append({
            "stem": "cluster_projection", "kind": "scatter", "data": embedding.head(20_000).copy(),
            "x": "dimension_1", "y": "dimension_2", "group": "cluster_label", "title": "Cluster projection",
            "x_label": "Projection dimension 1", "y_label": "Projection dimension 2",
            "commentary": "Points near one another have similar scaled indicator profiles in the two-dimensional projection. Separation supports descriptive typologies, but clusters are algorithmic labels—not natural classes, causal groups or performance ranks.",
        })

    forecast = next((table for name, table in result.tables.items() if "arima forecast" in name.casefold() and not table.empty), pd.DataFrame())
    if not forecast.empty and {"forecast_step", "forecast"}.issubset(forecast.columns):
        data = forecast.copy()
        data["series"] = "Forecast"
        specs.append({
            "stem": "arima_forecast", "kind": "line", "data": data,
            "x": "forecast_step", "y": "forecast", "group": "series", "title": "ARIMA forecast",
            "x_label": "Forecast step", "y_label": "Forecast value",
            "commentary": "The line is the model’s central forecast. The accompanying result table contains 95% intervals, which quantify model-based uncertainty conditional on the fitted order and absence of an unmodelled structural break.",
        })
    return specs[:12]


def _plotly_from_spec(spec: dict[str, Any]):
    import plotly.express as px
    data = spec["data"]
    kind = spec["kind"]
    if kind in {"barh", "signed_barh"}:
        colour = spec["x"] if kind == "signed_barh" else None
        fig = px.bar(data, x=spec["x"], y=spec["y"], orientation="h", color=colour, color_continuous_scale="RdBu" if colour else None, title=spec["title"])
    elif kind == "bar":
        fig = px.bar(data, x=spec["x"], y=spec["y"], title=spec["title"])
    elif kind == "line":
        fig = px.line(data, x=spec["x"], y=spec["y"], color=spec.get("group"), markers=True, title=spec["title"])
    elif kind == "scatter":
        fig = px.scatter(data, x=spec["x"], y=spec["y"], color=spec.get("group"), title=spec["title"], opacity=.75)
    elif kind == "hist":
        fig = px.histogram(data, x=spec["x"], nbins=60, marginal="box", title=spec["title"])
    else:
        error_x = None
        error_x_minus = None
        if {"ci_95_low", "ci_95_high"}.issubset(data.columns):
            data = data.copy()
            data["error_plus"] = data.ci_95_high - data.coefficient
            data["error_minus"] = data.coefficient - data.ci_95_low
            error_x, error_x_minus = "error_plus", "error_minus"
        fig = px.scatter(data, x=spec["x"], y=spec["y"], error_x=error_x, error_x_minus=error_x_minus, title=spec["title"])
        fig.add_vline(x=0, line_dash="dash", line_color="#555555")
    fig.update_layout(template="plotly_white", font=dict(family="Arial", size=13, color="#111111"), paper_bgcolor="white", plot_bgcolor="white", margin=dict(l=80, r=35, t=70, b=60))
    fig.update_xaxes(title=spec.get("x_label", spec.get("x", "")), gridcolor="#E5E7EB")
    fig.update_yaxes(title=spec.get("y_label", spec.get("y", "")), gridcolor="#F1F3F5")
    return fig


def _matplotlib_from_spec(spec: dict[str, Any], monochrome: bool):
    from matplotlib import pyplot as plt
    data = spec["data"]
    kind = spec["kind"]
    main = "#555555" if monochrome else "#155B8A"
    accent = "#111111" if monochrome else "#C45A36"
    fig, ax = plt.subplots(figsize=(9.4, 6.2), constrained_layout=True)
    if kind in {"barh", "signed_barh"}:
        colours = [accent if float(v) < 0 else main for v in data[spec["x"]]] if kind == "signed_barh" else main
        ax.barh(data[spec["y"]].astype(str), data[spec["x"]], color=colours, edgecolor="#222222", linewidth=.4, hatch="//" if monochrome else "")
        if kind == "signed_barh": ax.axvline(0, color="#333333", linewidth=.9)
    elif kind == "bar":
        ax.bar(data[spec["x"]].astype(str), data[spec["y"]], color=main, edgecolor="#222222", linewidth=.4, hatch="//" if monochrome else "")
    elif kind == "line":
        for index, (label, part) in enumerate(data.groupby(spec.get("group", "measure"), dropna=False)):
            ax.plot(part[spec["x"]], part[spec["y"]], marker=["o", "s", "^", "D"][index % 4], linestyle=["-", "--", "-.", ":"][index % 4], color=[main, accent, "#6F7D3C", "#9A5C83"][index % 4] if not monochrome else ["#111111", "#555555", "#888888", "#BBBBBB"][index % 4], label=str(label))
        ax.legend(frameon=False, fontsize=8)
    elif kind == "scatter":
        for index, (label, part) in enumerate(data.groupby(spec.get("group", "cluster_label"), dropna=False)):
            colour = [main, accent, "#6F7D3C", "#9A5C83", "#D89B2B", "#3B7D6E"][index % 6] if not monochrome else ["#111111", "#444444", "#777777", "#999999", "#BBBBBB", "#DDDDDD"][index % 6]
            ax.scatter(part[spec["x"]], part[spec["y"]], s=22, alpha=.7, color=colour, label=str(label), marker=["o", "s", "^", "D", "v", "P"][index % 6])
        ax.legend(frameon=False, fontsize=8, ncol=2)
    elif kind == "hist":
        ax.hist(data[spec["x"]], bins="auto", color=main, edgecolor="white", linewidth=.5, hatch="//" if monochrome else "")
        ax.axvline(0, color=accent, linestyle="--", linewidth=1.2)
    else:
        y_pos = np.arange(len(data))
        xerr = None
        if {"ci_95_low", "ci_95_high"}.issubset(data.columns):
            xerr = np.vstack([data.coefficient - data.ci_95_low, data.ci_95_high - data.coefficient])
        ax.errorbar(data[spec["x"]], y_pos, xerr=xerr, fmt="o", color=main, ecolor=accent, capsize=2.5)
        ax.set_yticks(y_pos, data[spec["y"]].astype(str))
        ax.axvline(0, color="#333333", linestyle="--", linewidth=.9)
    ax.set_title(spec["title"], loc="left", fontsize=15, fontweight="bold")
    ax.set_xlabel(spec.get("x_label", spec.get("x", "")))
    ax.set_ylabel(spec.get("y_label", "" if kind in {"barh", "signed_barh", "coefficient"} else spec.get("y", "")))
    ax.grid(axis="x" if kind in {"barh", "signed_barh", "coefficient"} else "y", color="#D9D9D9", linewidth=.55)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(labelsize=8.5)
    fig.text(.01, .005, "Source: saved Research Chair analytical scope. Interpret within the declared protocol and limitations.", fontsize=7.5, color="#444444")
    return fig


def _write_chair_figures(archive: zipfile.ZipFile, scoped_data: pd.DataFrame, result: ProtocolResult) -> pd.DataFrame:
    """Write interactive HTML plus colour/B&W 600-dpi and vector figure sets."""
    from matplotlib import pyplot as plt
    rows: list[dict[str, Any]] = []
    for number, spec in enumerate(_chair_figure_specs(scoped_data, result), start=1):
        stem = f"figure_{number:02d}_{spec['stem']}"
        interactive = _plotly_from_spec(spec)
        archive.writestr(f"figures/interactive/{stem}.html", interactive.to_html(full_html=True, include_plotlyjs=True).encode("utf-8"))
        archive.writestr(f"figures/data/{stem}.csv", spec["data"].to_csv(index=False).encode("utf-8-sig"))
        for style, monochrome in (("colour", False), ("black_white", True)):
            figure = _matplotlib_from_spec(spec, monochrome)
            for fmt in ("png", "svg", "pdf"):
                payload = io.BytesIO()
                figure.savefig(payload, format=fmt, dpi=600 if fmt == "png" else None, bbox_inches="tight", facecolor="white")
                archive.writestr(f"figures/{style}/{stem}.{fmt}", payload.getvalue())
            plt.close(figure)
        rows.append({
            "figure_number": number,
            "title": spec["title"],
            "interactive_html": f"figures/interactive/{stem}.html",
            "publication_files": f"figures/colour/{stem}.[png|svg|pdf]; figures/black_white/{stem}.[png|svg|pdf]",
            "data_file": f"figures/data/{stem}.csv",
            "what_it_shows_and_means": spec["commentary"],
            "resolution": "PNG 600 dpi; SVG/PDF vector; standalone Plotly HTML",
        })
    return pd.DataFrame(rows)


def chair_interactive_figures(scoped_data: pd.DataFrame, result: ProtocolResult) -> list[tuple[str, Any, str]]:
    """Return the same source-grounded figures used by the downloadable bundle."""
    return [
        (spec["title"], _plotly_from_spec(spec), spec["commentary"])
        for spec in _chair_figure_specs(scoped_data, result)
    ]


def research_bundle(
    scoped_data: pd.DataFrame,
    protocol: dict[str, Any],
    result: ProtocolResult,
    evidence: pd.DataFrame,
    blueprint: str,
    question_answers: str = "",
) -> bytes:
    charts = ["Selected longitudinal evidence chart"] if not result.tables.get("Longitudinal results", pd.DataFrame()).empty else []
    guide = build_output_guide(result.tables, charts)
    guide_markdown = output_guide_markdown(guide)
    prompts = prompt_library()
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "README.txt",
            "Free/offline Research Command Chair bundle. Every result is conditional on the saved scope and protocol.\n\n"
            "START HERE:\n"
            "1. Open OUTPUT_INTERPRETATION_GUIDE.md for a plain-language explanation of every table and chart.\n"
            "2. Open research_command_results.xlsx; the Output guide sheet explains purpose, reading, meaning, limitations and next steps.\n"
            "3. Open QUESTION_ANSWERS.md for the feasibility verdict and genuine prose answer to every selected question.\n"
            "4. Open figures/FIGURE_INDEX.csv, then use the standalone interactive HTML or publication-ready 600-dpi/vector files.\n"
            "5. Use makryvelios_prompt_library.md for copy-ready questions across all dashboard modules.\n"
            "6. Check protocol.json before reporting results, because it records the exact variables, filters, equation and settings.\n"
            "7. Statistical association, prediction and simulation do not automatically establish causality.\n"
        )
        archive.writestr("protocol.json", json.dumps(protocol, ensure_ascii=False, indent=2, default=str))
        archive.writestr("paper_blueprint.md", blueprint)
        archive.writestr("paper_blueprint.docx", docx_bytes(blueprint))
        archive.writestr("QUESTION_ANSWERS.md", question_answers or "No question batch was run before this export.\n")
        archive.writestr("QUESTION_ANSWERS.docx", docx_bytes(question_answers or "# Question answers\n\nNo question batch was run before this export."))
        archive.writestr("OUTPUT_INTERPRETATION_GUIDE.md", guide_markdown)
        archive.writestr("output_interpretation_guide.csv", guide.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("makryvelios_prompt_library.md", prompt_library_markdown(prompts))
        archive.writestr("makryvelios_prompt_library.csv", prompts.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("filtered_analytical_data.csv", scoped_data.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("selected_pdf_evidence.csv", evidence.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("selected_pdf_evidence.txt", evidence_text(evidence, 100_000))
        for name, table in result.tables.items():
            safe = re.sub(r"[^A-Za-z0-9_-]+", "_", name).strip("_").lower()
            archive.writestr(f"tables/{safe}.csv", table.to_csv(index=False).encode("utf-8-sig"))
        figure_index = _write_chair_figures(archive, scoped_data, result)
        archive.writestr("figures/FIGURE_INDEX.csv", figure_index.to_csv(index=False).encode("utf-8-sig"))
        figure_notes = "# Figure-by-figure interpretation\n\n" + "\n\n".join(
            f"## Figure {int(row.figure_number)}. {row.title}\n\n{row.what_it_shows_and_means}\n\nFiles: `{row.interactive_html}` and `{row.publication_files}`."
            for row in figure_index.itertuples(index=False)
        ) if not figure_index.empty else "# Figure-by-figure interpretation\n\nNo chart-compatible result table was generated.\n"
        archive.writestr("figures/FIGURE_COMMENTARY.md", figure_notes)

        html_figures = []
        for index, spec in enumerate(_chair_figure_specs(scoped_data, result)):
            html_figures.append(_plotly_from_spec(spec).to_html(full_html=False, include_plotlyjs=True if index == 0 else False))
        portable_report = f"""<!doctype html><html><head><meta charset=\"utf-8\"><meta name=\"viewport\" content=\"width=device-width,initial-scale=1\"><title>Research Chair analytical report</title><style>body{{font-family:Arial,sans-serif;max-width:1180px;margin:0 auto;padding:28px;color:#111;background:#fff;line-height:1.55}}h1,h2{{color:#133149}}.answer,.blueprint{{white-space:pre-wrap;background:#f7f9fb;border:1px solid #dce3e8;border-radius:10px;padding:18px}}.figure{{margin:26px 0;padding:14px;border:1px solid #e2e8ee;border-radius:12px}}</style></head><body><h1>Research Chair analytical report</h1><h2>Selected-question answers</h2><div class=\"answer\">{html.escape(question_answers or 'No question batch was run before this export.')}</div><h2>Interactive figures</h2>{''.join(f'<div class=\"figure\">{item}</div>' for item in html_figures)}<h2>Paper blueprint</h2><div class=\"blueprint\">{html.escape(blueprint)}</div></body></html>"""
        archive.writestr("interactive_research_report.html", portable_report.encode("utf-8"))
        workbook = io.BytesIO()
        with pd.ExcelWriter(workbook, engine="openpyxl") as writer:
            scoped_data.head(1_048_000).to_excel(writer, sheet_name="Filtered data", index=False)
            evidence.drop(columns=["text"], errors="ignore").to_excel(writer, sheet_name="PDF evidence index", index=False)
            guide.to_excel(writer, sheet_name="Output guide", index=False)
            figure_index.to_excel(writer, sheet_name="Figure guide", index=False)
            prompts.to_excel(writer, sheet_name="Prompt library", index=False)
            for number, (name, table) in enumerate(result.tables.items(), start=1):
                sheet = re.sub(r"[\\/*?:\[\]]", "_", name)[:25] + f"_{number}"
                table.head(1_048_000).to_excel(writer, sheet_name=sheet, index=False)
        archive.writestr("research_command_results.xlsx", workbook.getvalue())
    return buffer.getvalue()
